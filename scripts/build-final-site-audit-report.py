"""Build the COWIN MACHINE pre-production audit report.

This creates a local report only. It does not touch application source, public
assets, product data, external services, or deployment state.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "final-site-audit.docx"
SKILL_SCRIPTS = Path(
    r"D:\codex\.codex\plugins\cache\openai-primary-runtime\documents\26.812.11052\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
GOLD = "B26A00"
RISK = "9B1C1C"
HEADER_FILL = "E8EEF5"
CALLOUT = "F4F6F9"
WARN = "FFF4E5"
RISK_FILL = "FDECEC"
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell_or_paragraph, fill):
    if hasattr(cell_or_paragraph, "_tc"):
        properties = cell_or_paragraph._tc.get_or_add_tcPr()
    else:
        properties = cell_or_paragraph._p.get_or_add_pPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    element.set(qn("w:val"), "clear")
    properties.append(element)


def add_page_number(paragraph):
    run = paragraph.add_run()
    font(run, 8.5, MUTED)
    start = OxmlElement("w:fldChar")
    start.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([start, instruction, end])


def repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, size, color in (("Audit Small", 8.5, MUTED), ("Audit Caption", 9, MUTED)):
        if name not in doc.styles:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    font(header.add_run("COWIN MACHINE  |  Pre-Production Site Audit"), 8.5, MUTED, True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    font(footer.add_run("Internal release review  |  Page "), 8.5, MUTED)
    add_page_number(footer)


def paragraph(doc, text, *, style=None, color=None, size=None, bold_prefix=None, italic=False, align=None, fill=None):
    item = doc.add_paragraph(style=style)
    if align is not None:
        item.alignment = align
    if fill:
        shade(item, fill)
    if bold_prefix and text.startswith(bold_prefix):
        font(item.add_run(bold_prefix), size, color, True)
        font(item.add_run(text[len(bold_prefix):]), size, color, italic=italic)
    else:
        font(item.add_run(text), size, color, italic=italic)
    return item


def bullets(doc, values):
    for value in values:
        item = doc.add_paragraph(style="List Bullet")
        item.paragraph_format.space_after = Pt(4)
        item.paragraph_format.line_spacing = 1.25
        item.paragraph_format.left_indent = Inches(0.375)
        item.paragraph_format.first_line_indent = Inches(-0.188)
        font(item.add_run(value))


def table(doc, headers, rows, widths, risk_col=None):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.autofit = False
    for index, value in enumerate(headers):
        cell = result.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, HEADER_FILL)
        run = cell.paragraphs[0].add_run(value)
        font(run, 9, INK, True)
    repeat_header(result.rows[0])
    for row in rows:
        cells = result.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if risk_col is not None and index == risk_col and str(value).startswith("P0"):
                shade(cells[index], RISK_FILL)
            elif risk_col is not None and index == risk_col and str(value).startswith("P1"):
                shade(cells[index], WARN)
            text = cells[index].paragraphs[0]
            text.paragraph_format.space_after = Pt(2)
            font(text.add_run(str(value)), 8.5, INK, bold=index == 0)
    apply_table_geometry(result, [Inches(value) for value in widths], indent_dxa=TABLE_INDENT, cell_margins_dxa=CELL_MARGINS)
    return result


def issue(doc, code, heading, finding, impact, action, severity):
    doc.add_heading(f"{code} — {heading}", level=2)
    paragraph(doc, f"Finding: {finding}", bold_prefix="Finding: ")
    paragraph(doc, f"Impact: {impact}", bold_prefix="Impact: ")
    paragraph(doc, f"Required resolution: {action}", bold_prefix="Required resolution: ")


def build():
    doc = Document()
    setup(doc)

    # Editorial-cover opening pattern, with compact-reference-guide styles thereafter.
    for _ in range(6):
        paragraph(doc, "", size=4)
    paragraph(doc, "PRE-PRODUCTION AUDIT", color=GOLD, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "COWIN MACHINE", color=INK, size=30, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Website release readiness review", color=DARK_BLUE, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Product data, content, SEO, accessibility, performance, security and publishing controls", color=MUTED, size=10.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "14 August 2026", color=INK, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Release decision: BLOCKED — P0 and P1 remediation required", color=RISK, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, fill=RISK_FILL)
    doc.add_page_break()

    doc.add_heading("Executive summary", level=1)
    paragraph(doc, "The local application builds successfully and its core information architecture is implemented. It has six product categories, 202 product-detail routes, product navigation drawers, inquiry prefill, sitemap, robots, RSS, structured-data components, research records, and a draft-only content workflow. It is not ready for production: three P0 risks can expose unverified technical claims or unsupported image-rights claims, and six P1 issues affect product-page reliability, inquiry handling, security posture, or mobile conversion behavior.")
    paragraph(doc, "No production deployment, publication, data deletion, or Git push was performed during this review.", bold_prefix="No production deployment")

    doc.add_heading("Scope and methodology", level=1)
    paragraph(doc, "The review covered the requested product, fact, duplicate/canonical, image, inquiry, navigation, mobile, form, Word-report, content-workflow, SEO/GEO, schema, sitemap, robots, RSS, performance, accessibility, and security areas. Evidence was collected from local JSON/CSV datasets, application source and configuration, a local production build, a local route crawl, inquiry-route tests, and document-audit scripts.")
    paragraph(doc, "Limitations: This is a technical and evidence-availability audit. It cannot grant copyright permissions, validate undocumented supplier facts, guarantee search indexing, or replace a real-device browser acceptance test.", bold_prefix="Limitations: ", fill=CALLOUT)

    doc.add_heading("Evidence dashboard", level=1)
    table(doc, ["Measure", "Result", "Evidence"], [
        ("Current product records", "202 across six categories", "raw-product-inventory.json"),
        ("Canonical families", "135", "canonical-product-master.json"),
        ("Duplicate raw records", "87", "canonical master and issue report"),
        ("Verified-model families", "68", "canonical master"),
        ("Review/generic/misclassification families", "67", "canonical master"),
        ("Detail-page states", "68 full technical; 134 configuration review", "product-detail-content.json"),
        ("Route crawl", "218 seeds; 624 unique internal hrefs; all HTTP 200", "local production crawl"),
        ("News workflow", "6 eligible candidates; 3 pending review; 0 published", "content-state.json"),
    ], [1.65, 2.55, 2.3])

    doc.add_heading("Current records by category", level=2)
    table(doc, ["Category", "Current records"], [
        ("Compressed air equipment", "36"),
        ("Generator systems", "10"),
        ("Drilling equipment", "12"),
        ("Drilling consumables", "16"),
        ("Mobile lighting systems", "39"),
        ("Magnetic separators", "89"),
        ("Total", "202"),
    ], [4.7, 1.8])

    doc.add_heading("Priority summary", level=1)
    table(doc, ["Priority", "Open issues", "Release effect"], [
        ("P0", "3", "Blocks release: accuracy, schema, and image-rights risk"),
        ("P1", "6", "Resolve before any production recommendation"),
        ("P2", "6", "Improve before or immediately after controlled launch"),
        ("P3", "4", "Operational and growth maturity"),
    ], [1.0, 1.35, 4.15], risk_col=0)

    doc.add_heading("P0 — release blockers", level=1)
    issue(doc, "P0-01", "Review-state cards expose technical values", "134 product profiles are marked configuration review, but category cards still render key specifications. At least 41 review-state records include values beyond model identity and configuration.", "The public site can make technical commitments for records that its own product master says require confirmation.", "Render only non-technical identity/context and Request Configuration Review for every review-state record. Require source evidence before any technical value is displayed, then re-crawl every category.", "P0")
    issue(doc, "P0-02", "Product schema uses specifications without independent evidence", "The 68 verified-model families have Product JSON-LD and technical values, but their recorded evidence is current-site only. No attached owner-approved datasheet, manufacturer document, or equivalent independent source supports the field values.", "Unsupported technical values can be surfaced in page content and machine-readable schema.", "Attach field-level source evidence or remove the affected page content and schema additional properties. Do not output offers, ratings, warranties, performance, capacity, or certification claims without evidence.", "P0")
    issue(doc, "P0-03", "Image authorization is asserted but not auditable", "All 202 product detail profiles label images as authorized, but the image-asset manifest has no per-asset source, permission, watermark review, or approval record. The mag-series-080-080 route has no configured hero image.", "The site can make an unsupported rights claim or publish an old-brand, watermarked, or unlicensed image.", "Create a versioned rights manifest for every public asset and hide authorization labels until supported. Replace or omit undocumented assets and provide a compliant image/no-image state for the missing hero.", "P0")

    doc.add_heading("P1 — resolve before deployment recommendation", level=1)
    issue(doc, "P1-01", "Visible undefined product overview", "Review-state pages render copy beginning with 'undefined is retained as a …'; it was reproduced locally on a magnetic-separator review route.", "Broken public product copy reduces trust and can be indexed.", "Use a safe status template that never interpolates an absent model field; route-test all 134 review pages.", "P1")
    issue(doc, "P1-02", "Unsupported 24-hour inquiry promise", "The inquiry API success response promises a 24-hour reply, but no configured response SLA, email service, CRM, or database delivery workflow exists.", "The site makes an operational commitment that cannot be proven or monitored.", "Use a non-committal receipt message, or implement and owner-approve a real response workflow and service commitment.", "P1")
    issue(doc, "P1-03", "No production-safe inquiry persistence or rate limiting", "The form validation and honeypot work locally, but submissions are not durably stored and the endpoint lacks deployable rate control. Local file storage intentionally blocks writes on Vercel.", "Spam and repeated requests can reach the handler; valid inquiries cannot be reliably routed to sales.", "Add server-side rate limiting and a production-safe storage/CRM adapter configured solely via environment variables. Test success, retry, and failure states.", "P1")
    issue(doc, "P1-04", "Security headers are absent", "No CSP, Permissions-Policy, Referrer-Policy, X-Content-Type-Options, frame-ancestor policy, or HTTPS-edge HSTS configuration was found.", "A public lead-generation site has a weaker default browser-security posture.", "Add and validate a tested policy on preview. Configure HSTS at the HTTPS edge only after redirect/domain validation.", "P1")
    issue(doc, "P1-05", "Mobile quote CTA appears after navigation", "The mobile menu presents Products and other links before Get a Quote, contrary to the stated conversion requirement.", "Primary conversion is less discoverable and the mobile requirement is not met.", "Move Request a Quote to the top of the mobile drawer and manually test all breakpoints.", "P1")
    issue(doc, "P1-06", "No independent visual acceptance evidence", "Code review indicates intended focus, escape, overlay, scroll-lock, reduced-motion, and responsive behavior, but Chrome headless visual capture failed in this environment during GPU startup.", "Real-browser layout, focus visibility, and touch interaction remain unverified.", "Record manual acceptance on Chrome, Safari/iOS, and Android Chrome at 320 px, 768 px, 1024 px, and desktop widths.", "P1")

    doc.add_heading("P2 — completeness and operational improvements", level=1)
    bullets(doc, [
        "Use true content-update timestamps for sitemap lastmod values instead of the generation date.",
        "Explicitly disallow internal operational paths and non-public API routes in robots policy.",
        "Validate the first approved news article end-to-end: RSS, Article JSON-LD, visible FAQ, canonical, and sitemap inclusion.",
        "Re-run all quality gates for the three pending-review drafts created before the latest quality-gate changes.",
        "Render and visually inspect the three existing Word research reports on an environment with a LibreOffice-compatible renderer.",
        "Add link, accessibility, schema, review-state, drawer keyboard-flow, and 320 px smoke tests to CI.",
    ])

    doc.add_heading("P3 — growth and governance opportunities", level=1)
    bullets(doc, [
        "Add a Search Console adapter only after verified credentials and ownership are configured; report unknown or pending states honestly.",
        "Create a controlled asset-review workflow for alt text, crops, AI-generated assets, stock licences, and supplier approvals.",
        "Define product-owner evidence-refresh intervals and a supplier-document intake checklist.",
        "Use preview deployments and versioned content snapshots for every product, schema, sitemap, and news release.",
    ])

    doc.add_heading("Confirmed controls and positive findings", level=1)
    bullets(doc, [
        "The reviewed public application scope contained no prohibited competitor/legacy company names, logos, addresses, telephone numbers, tracking identifiers, or external image hotlinks. The supplied COWIN MACHINE contact email includes cowinmagnet.com as its domain, but it is not displayed as a competing brand.",
        "Organization schema is limited to the supplied identity fields, and local canonical spot checks point to COWIN MACHINE route paths.",
        "Configuration-review detail pages omit Product JSON-LD; this is correct, though their category-card leak and copy defect must still be fixed.",
        "The local route crawl found 624 unique internal href targets; every tested target returned HTTP 200. The build generated all 202 product detail routes.",
        "The inquiry route handled valid, validation-error, and honeypot paths locally without exposing secrets.",
        "Products and site menu drawers are separate mutually exclusive client components; mobile Products uses an accordion.",
        "Content automation defaults to draft mode with auto-publish disabled. Current state contains three pending-review items and zero published articles.",
        "Local lint, TypeScript checking, content-automation validation, and production build completed successfully.",
    ])

    doc.add_heading("News, SEO, GEO, and documents", level=1)
    paragraph(doc, "The news system is an internal source-candidate and original-draft workflow, not a published article corpus. It tracks source quality, date eligibility, originality risk, internal-link requirements, and draft/publish controls. It is a useful operating baseline, not proof that future articles will automatically be original or accurate. Human approval must remain in place for source facts, product evidence, image rights, and similarity checks.")
    paragraph(doc, "Sitemap, robots, RSS, canonical, Open Graph, Product, BreadcrumbList, Article, and conditional FAQ structured-data components are implemented. They improve discovery and machine readability but cannot guarantee inclusion in Google, AI systems, or any other search surface.")
    paragraph(doc, "The product research, industry/application research, and 90-day news plan Word reports passed their structural heading/accessibility scripts with zero reported findings. They are internal evidence baselines, not legal authorization for images, data sheets, certifications, or unverified product parameters.")

    doc.add_heading("Release sequence and rollback", level=1)
    bullets(doc, [
        "Resolve and regression-test every P0 and P1 item before preparing a release candidate.",
        "Create a product-data snapshot, field-evidence bundle, image-rights manifest, and owner sign-off record.",
        "Deploy only to preview first, then test forms, routes, drawers, schema, sitemap, RSS, security headers, and mobile layouts.",
        "Tag the approved Git commit and retain the prior deployment ID plus the prior content-state/evidence snapshot.",
        "For rollback, promote the prior deployment, restore the prior approved content state, keep draft mode and auto-publish disabled, then re-check public discovery files.",
    ])

    doc.add_heading("Deployment decision", level=1)
    paragraph(doc, "DO NOT DEPLOY. P0 and P1 issues are open. Recommend production deployment only after the required evidence, technical-content gates, inquiry delivery and rate controls, security headers, mobile CTA ordering, and real-device acceptance checks are complete and documented.", color=RISK, size=11.5, fill=RISK_FILL)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
