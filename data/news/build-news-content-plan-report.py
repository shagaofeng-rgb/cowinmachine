"""Build the local-only 90-day industry news and original-content plan."""

from __future__ import annotations

import json
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
NEWS = ROOT / "data" / "news"
REPORT = ROOT / "docs" / "COWIN-MACHINE-90-Day-News-and-Content-Plan.docx"
SKILL_SCRIPTS = Path(r"D:\codex\.codex\plugins\cache\openai-primary-runtime\documents\26.812.11052\skills\documents\scripts")
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402

NAVY, BLUE, ORANGE, INK, MUTED = "0B1F33", "1B4D72", "F58A1F", "172033", "5B6878"
HEADER, NOTE, RISK = "EAF0F5", "F5F7FA", "FFF4E5"
WINDOW = "16 May 2026 to 14 August 2026"


def compact(value: object) -> str:
    return " ".join(str(value or "").split())


def font(run, size=None, color=None, bold=None):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade(cell, value):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), value)
    node.set(qn("w:val"), "clear")
    props.append(node)


def setup(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        header = section.header.paragraphs[0]
        header.text = ""
        font(header.add_run("COWIN MACHINE  |  NEWS RESEARCH & ORIGINAL CONTENT PLAN"), 8.5, MUTED, True)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.text = ""
        font(footer.add_run("Local planning document  |  Page "), 8.5, MUTED)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 17, 8),
        ("Heading 2", 12.5, BLUE, 13, 6),
        ("Heading 3", 11, NAVY, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_p(doc, text, style=None, color=None, size=None, bold_prefix=None, fill=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if fill:
        props = p._p.get_or_add_pPr()
        node = OxmlElement("w:shd")
        node.set(qn("w:fill"), fill)
        node.set(qn("w:val"), "clear")
        props.append(node)
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), size, color, True)
        font(p.add_run(text[len(bold_prefix):]), size, color)
    else:
        font(p.add_run(text), size, color)
    return p


def h(doc, text, level=1):
    return add_p(doc, text, style=f"Heading {level}")


def bullets(doc, values):
    for value in values:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(value), 9.7)


def table(doc, headers, rows, widths, size=7.8):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    header = result.rows[0]
    header_props = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_props.append(repeat)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, HEADER)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(value), size, INK, True)
    for row in rows:
        cells = result.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            font(p.add_run(compact(value)), size, INK)
    apply_table_geometry(result, widths, table_width_dxa=sum(widths), indent_dxa=80, cell_margins_dxa={"top":70, "bottom":70, "start":90, "end":90})
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def load(filename):
    return json.loads((NEWS / filename).read_text(encoding="utf-8"))


def queue_category(item):
    parts = item["primaryProductUrl"].strip("/").split("/")
    return parts[1] if len(parts) > 1 else "unmapped"


def main():
    candidates = load("source-candidates.json")["candidates"]
    approved = load("approved-sources.json")["sources"]
    queue = load("article-queue.json")["articles"]
    rejected = load("rejected-topics.json")["rejected"]
    eligible = [item for item in candidates if item["eligibleForArticle"]]
    ready = [item for item in queue if item["status"] == "research-ready"]
    coverage = Counter(queue_category(item) for item in queue)

    doc = Document()
    setup(doc)
    add_p(doc, "90-DAY INDUSTRY NEWS RESEARCH", size=10.5, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "COWIN MACHINE", size=28, color=NAVY, bold_prefix="COWIN MACHINE", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "News Research and Original Content Plan", size=18, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "A local review system for source screening, original drafts, quality checks and a non-publishing article queue", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    add_p(doc, "Research window: " + WINDOW, size=10, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Generated 14 August 2026  |  No article has been published or scheduled for publication.", size=9.2, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    h(doc, "1. Executive summary")
    add_p(doc, f"This plan establishes a controlled, local-only workflow for collecting high-confidence industry-news candidates and turning only approved source combinations into original COWIN MACHINE drafts. The initial pass records {len(candidates)} candidates: {len(eligible)} are recent and eligible for research, while {len(rejected)} are deliberately rejected. The editorial queue contains {len(queue)} original topic briefs across six product categories; {len(ready)} meet the initial two-source research threshold for drafting, not publication.")
    add_p(doc, "Publication control: ", bold_prefix="Publication control: ", fill=RISK)
    add_p(doc, "No queue item can be published automatically. Publication remains blocked until the owner approves the factual scope, the product-family evidence, the source links, the copy, metadata, internal links and any image rights.", fill=RISK)

    h(doc, "2. Scope, monitoring window and methodology")
    add_p(doc, "The system aims to continuously collect relevant, citable global industry-news candidates from the preceding 90 days. It does not claim to collect all global news. This first research window runs from 16 May 2026 through 14 August 2026.")
    bullets(doc, [
        "The candidate register preserves source URL, published date, discovery date, category and industry mapping, concise fact summaries, quality level, image-rights status and originality risk.",
        "A recent source can add industry context; it does not validate a COWIN MACHINE model, performance, certification, price, delivery term or customer case.",
        "Original drafts summarize a limited fact from each source, then provide independent operational and selection context. They do not translate, rewrite or reproduce source articles.",
        "Third-party imagery, logos, PDFs, certificates, customer stories and hot-linked images are excluded from the workflow.",
    ])

    h(doc, "3. Current website relevance and content gate")
    add_p(doc, "The plan maps article briefs to existing COWIN MACHINE product-family routes only as internal navigation targets. Technical claims remain subject to the existing product evidence controls. A generic, duplicate, missing-specification or misclassified product family cannot be promoted by a news article.")
    table(doc, ["Gate", "Requirement", "Outcome if not met"], [
        ["Product evidence", "Verified family or approved category-level technical framing", "Use a configuration-review CTA; do not claim model performance."],
        ["Recent reporting", "At least two independent, eligible 90-day sources for a publishable research brief", "Keep the topic research-pending."],
        ["Originality", "New H1, summary, structure, FAQ, examples and internal-link context", "Rewrite; do not use mechanical paraphrase."],
        ["Rights", "No third-party protected image, PDF, logo or unapproved case asset", "Use no image or an approved new asset."],
        ["Owner review", "Approval of facts, claims, CTA, metadata and destination URLs", "Remain local-draft-only."],
    ], [1750, 4750, 2860], 8.2)

    h(doc, "4. Candidate-news inventory")
    add_p(doc, "The records below are research inputs, not COWIN MACHINE endorsements. Source summaries are intentionally concise and no external text or visual asset is copied.")
    table(doc, ["ID", "Source", "Date", "Category / industry", "Eligibility", "Research decision"], [
        [item["id"], item["sourceName"], item["publishedAt"], "; ".join(item["productCategories"] + item["industries"]), "Eligible" if item["eligibleForArticle"] else "Rejected", "Use as context" if item["eligibleForArticle"] else item.get("rejectionReason", "Rejected")]
        for item in candidates
    ], [800, 1800, 900, 2650, 900, 2310], 7.1)

    h(doc, "5. Approved source directory")
    add_p(doc, "This directory includes primary announcements and durable government, association or technical references. The directory is a monitoring and fact-checking list; it is not a license to republish source material.")
    table(doc, ["Source", "Type", "Topics", "Use in workflow", "Restrictions"], [
        [item["name"], item["quality"], "; ".join(item["coverage"]), item["use"], "Original source review required; no asset or copy reuse."]
        for item in approved
    ], [1750, 1250, 1750, 2800, 1810], 7.0)

    h(doc, "6. Original-content production system")
    add_p(doc, "A local draft generator accepts only research-ready queue items with two independent eligible sources. Each generated English draft targets approximately 1,200–1,800 words and is stored locally for review. It is never pushed to a route, CMS, deployment or social channel.")
    table(doc, ["Module", "Draft requirement"], [
        ["Answer-first introduction", "50–70 words that answer the operational question without product promises."],
        ["Technical and application context", "Equipment overview, workflow, pain points, selection considerations and practical benefits."],
        ["Recent developments", "Short, attributed fact summaries with source links; no long quotations or copied paragraphs."],
        ["Buyer guidance", "A concrete next-evaluation list, relevant FAQ and a configuration-review CTA."],
        ["SEO / GEO checks", "Unique H1, title, description, canonical plan, Article/Breadcrumb/FAQ schema plan, and at least three internal links."],
    ], [2600, 6760], 8.0)

    h(doc, "7. 36-topic original content calendar")
    add_p(doc, "Coverage: six briefs per category. Only the three research-ready topics have local review drafts; the other 33 remain research-pending until their source requirements are met.")
    table(doc, ["ID", "Original working title", "Product category", "Industry", "Priority", "Status"], [
        [item["id"], item["title"], queue_category(item), item["industry"], item["priority"], item["status"]]
        for item in queue
    ], [700, 3150, 1600, 1450, 850, 1610], 7.0)
    h(doc, "8. Category coverage")
    table(doc, ["Product category", "Topic briefs", "Editorial focus"], [
        [category, str(count), "Industry application, buyer selection and source-backed current context"]
        for category, count in sorted(coverage.items())
    ], [3200, 1500, 4660], 8.2)

    h(doc, "9. Rejections and quality controls")
    add_p(doc, "Rejected candidates remain in a dedicated register so that the research trail is auditable and the same low-quality source is not accidentally promoted later.")
    table(doc, ["Candidate", "Source", "Reason for rejection", "Control"], [
        [item["id"], item["sourceName"], item.get("reason", "Not eligible"), "Not available to the draft generator"]
        for item in rejected
    ], [1500, 2100, 3900, 1860], 7.5)
    bullets(doc, [
        "Reject sources without accountable authorship, date, provenance or usable evidence.",
        "Reject any route that would make another company’s project appear to be a COWIN MACHINE customer case.",
        "Block titles, descriptions, H2s, FAQ answers and image alt text that are mechanically similar to source material or another planned article.",
        "Do not invent authors, experts, tests, product measurements, project outcomes, warranties, certifications or market claims.",
    ])

    h(doc, "10. Publishing queue and next actions")
    add_p(doc, "The system’s release state is intentionally local-only. `published-history.json` is empty and serves as the future duplicate-prevention ledger after a separately authorized publication workflow exists.")
    bullets(doc, [
        "Review the three local drafts for product-family accuracy, legal/brand safety and source context before any implementation request.",
        "Continue monitoring the approved source directory and add new candidates only after checking the original page, publication date and source quality.",
        "For each research-pending topic, collect at least two independent, in-window sources or reframe it as an evergreen technical resource outside this news workflow.",
        "Before future publication, run copy-similarity, claims, link, schema, canonical, accessibility and image-rights checks; then obtain explicit production approval.",
    ])

    doc.core_properties.title = "COWIN MACHINE 90-Day News and Content Plan"
    doc.core_properties.subject = "Local-only research and original-content planning"
    doc.core_properties.author = "COWIN MACHINE content workflow"
    doc.save(REPORT)
    print(json.dumps({"report": str(REPORT), "candidates": len(candidates), "eligibleRecent": len(eligible), "queue": len(queue), "researchReady": len(ready), "rejected": len(rejected)}, indent=2))


if __name__ == "__main__":
    main()
