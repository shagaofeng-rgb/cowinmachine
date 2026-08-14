"""Build a research-only global technical and application library.

This generator consumes the approved local product-audit outputs and writes
research dossiers, evidence registers, matrices, and a Word report. It does
not alter application source code, product pages, public assets, or deployment
configuration.
"""

from __future__ import annotations

import csv
import json
import sys
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
AUDIT_DIR = ROOT / "data" / "product-audit"
RESEARCH_DIR = ROOT / "data" / "research"
DOSSIER_DIR = RESEARCH_DIR / "product-dossiers"
DOCS_RESEARCH_DIR = ROOT / "docs" / "research"
REPORT_PATH = ROOT / "docs" / "COWIN-MACHINE-Industry-and-Application-Research.docx"
SKILL_SCRIPTS = Path(r"D:\codex\.codex\plugins\cache\openai-primary-runtime\documents\26.812.11052\skills\documents\scripts")
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402

ACCESS_DATE = "2026-08-14"
SITE = "https://cowinmachine.com"
BLUE, DARK, INK, MUTED = "2E74B5", "1F4D78", "0B2545", "667085"
HEADER, NOTE, RISK = "E8EEF5", "F4F6F9", "FFF4E5"
CATEGORY_LABELS = {
    "compressed-air-equipment": "Air Compressors",
    "generator-systems": "Generator Systems",
    "drilling-equipment": "Drilling Rigs",
    "drilling-consumables": "Drilling Tools & Consumables",
    "mobile-lighting-systems": "Solar & Mobile Light Towers",
    "magnetic-separators": "Magnetic Separators",
}

# These records were individually opened and reviewed during the research pass.
# They support generic technical explanations only. No supplier imagery, PDFs,
# brand copy, model data, certification, or commercial claims are reused.
EXTERNAL_SOURCES = [
    {
        "sourceId": "SRC-CA-01", "categories": ["compressed-air-equipment"],
        "sourceName": "Compressed Air Systems", "organization": "U.S. Department of Energy",
        "sourceType": "government technical guidance", "url": "https://www.energy.gov/cmei/ito/compressed-air-systems",
        "evidenceSummary": "Identifies air quality, condensate removal, leak control, pressure stabilization, controls, storage and preventive maintenance as compressed-air system topics.",
        "confidence": "high", "useRestriction": "General system guidance only; no manufacturer product data or assets used.",
    },
    {
        "sourceId": "SRC-GEN-01", "categories": ["generator-systems"],
        "sourceName": "Generator Ratings: Prime, Standby Power and More", "organization": "Caterpillar",
        "sourceType": "manufacturer technical article", "url": "https://www.cat.com/en_GB/by-industry/electric-power/Articles/White-papers/understanding-generator-set-ratings.html",
        "evidenceSummary": "Explains duty concepts including emergency standby, prime, limited-time running and continuous operation in relation to ISO 8528-1 terminology.",
        "confidence": "high", "useRestriction": "General rating definitions only; do not reuse Caterpillar models, data, graphics, certification or wording.",
    },
    {
        "sourceId": "SRC-GEN-02", "categories": ["generator-systems"],
        "sourceName": "Commercial Generator Sizing", "organization": "Caterpillar",
        "sourceType": "manufacturer technical guide", "url": "https://www.cat.com/en_US/by-industry/electric-power/electric-power-resources/genset-sizing.html",
        "evidenceSummary": "Describes considering running and starting loads, voltage, frequency, phase, environmental derating, motor inrush and load steps during sizing.",
        "confidence": "high", "useRestriction": "Selection-factor reference only; no calculator output, product data, images or copy reused.",
    },
    {
        "sourceId": "SRC-DR-01", "categories": ["drilling-equipment", "drilling-consumables"],
        "sourceName": "Quarrying", "organization": "Epiroc",
        "sourceType": "manufacturer application guidance", "url": "https://www.epiroc.com/en-ph/applications/construction/quarrying-and-surface-construction/quarrying",
        "evidenceSummary": "Links drilling pattern, geology, hole diameter and fragmentation requirements in quarry and surface-construction work, and notes DTH and tophammer use.",
        "confidence": "medium", "useRestriction": "General application context only; no Epiroc model, branding, images, performance or customer content reused.",
    },
    {
        "sourceId": "SRC-DR-02", "categories": ["drilling-equipment", "drilling-consumables"],
        "sourceName": "DTH Equipment Product Catalogue", "organization": "Epiroc",
        "sourceType": "manufacturer technical catalogue", "url": "https://www.epiroc.com/content/dam/epiroc/rock-drilling-tools/documents/Epiroc%20DTH%20product%20catalog.pdf",
        "evidenceSummary": "Supports the generic principle that hammer, bit and hole geometry must be matched so cuttings can be evacuated; working pressure, interface and hole diameter are selection factors.",
        "confidence": "medium", "useRestriction": "General DTH mechanism only. A matching-looking model code does not establish COWIN MACHINE manufacturer, compatibility or specification.",
    },
    {
        "sourceId": "SRC-DR-03", "categories": ["drilling-equipment", "drilling-consumables", "compressed-air-equipment"],
        "sourceName": "Best Practices for Dust Control in Surface Drills", "organization": "NIOSH / CDC",
        "sourceType": "government safety guidance", "url": "https://www.cdc.gov/niosh/engcontrols/ecd/detail72.html",
        "evidenceSummary": "Describes drill dust generated by air flushing, dry or water-based control options, and the need to inspect shrouds, seals, filters and ducting.",
        "confidence": "high", "useRestriction": "Safety context only; site-specific controls must be engineered and comply with local law.",
    },
    {
        "sourceId": "SRC-DR-04", "categories": ["drilling-equipment", "drilling-consumables"],
        "sourceName": "Equipment Condition", "organization": "Occupational Safety and Health Administration",
        "sourceType": "government safety guidance", "url": "https://www.osha.gov/etools/oil-and-gas/general-safety/equipment-condition",
        "evidenceSummary": "Emphasizes correct pressure-rated equipment selection, matching parts, pre-start inspection, documented maintenance and manufacturer procedures.",
        "confidence": "high", "useRestriction": "Safety guidance only; it is not a product approval or operating manual.",
    },
    {
        "sourceId": "SRC-LT-01", "categories": ["mobile-lighting-systems"],
        "sourceName": "Solar Energy and Storage Basics", "organization": "U.S. Department of Energy / NREL",
        "sourceType": "government technical guidance", "url": "https://www.energy.gov/cmei/systems/solar-integration-solar-energy-and-storage-basics",
        "evidenceSummary": "Explains that solar production varies with time, weather and shading; storage shifts energy use in time and energy capacity differs from power capacity.",
        "confidence": "high", "useRestriction": "General solar-and-storage principles only; no runtime or capacity is implied for any COWIN MACHINE product.",
    },
    {
        "sourceId": "SRC-LT-02", "categories": ["mobile-lighting-systems"],
        "sourceName": "Hybrid LT-Series", "organization": "Allmand",
        "sourceType": "manufacturer product documentation", "url": "https://www.allmand.com/products/light-towers/hybrid-lt-series/",
        "evidenceSummary": "Confirms a generic mobile tower configuration can combine luminaires, mast, stabilizers, trailer, control system and diesel/battery operating modes.",
        "confidence": "medium", "useRestriction": "Configuration reference only. No Allmand values, text, graphics, manual content, model names or commercial claims are reused.",
    },
    {
        "sourceId": "SRC-LT-03", "categories": ["mobile-lighting-systems"],
        "sourceName": "Construction Illumination", "organization": "Occupational Safety and Health Administration",
        "sourceType": "government safety regulation", "url": "https://www.osha.gov/laws-regs/regulations/standardnumber/1926/1926.26",
        "evidenceSummary": "States that active construction areas, access routes, offices, shops and storage areas require natural or artificial illumination.",
        "confidence": "high", "useRestriction": "Use as a safety context only; local requirements and lighting design must be confirmed for the installation.",
    },
    {
        "sourceId": "SRC-LT-04", "categories": ["mobile-lighting-systems"],
        "sourceName": "Battery Energy Storage System Procurement Checklist", "organization": "U.S. Department of Energy",
        "sourceType": "government procurement guidance", "url": "https://www.energy.gov/cmei/femp/articles/battery-energy-storage-system-procurement-checklist",
        "evidenceSummary": "Shows that battery-storage procurement should use technical specifications and project checklists rather than an assumed generic configuration.",
        "confidence": "high", "useRestriction": "Does not validate battery chemistry, capacity, runtime or safety certification for any product.",
    },
    {
        "sourceId": "SRC-MS-01", "categories": ["magnetic-separators"],
        "sourceName": "Wet Drum Separators", "organization": "Eriez",
        "sourceType": "manufacturer technical page", "url": "https://www.eriez.com/Products/Magnetic-Separation/Magnetic-Processing-Concentration/Magnetic-Separation-Solutions-for-Wet-Minerals-Processing/Wet-Drum-Separators",
        "evidenceSummary": "Explains a generic wet-drum process in which slurry enters a tank, magnetic material is retained on a rotating drum and nonmagnetics exit separately; tank configuration influences the duty.",
        "confidence": "medium", "useRestriction": "Generic working principle only; do not reuse Eriez performance, components, images or model statements.",
    },
    {
        "sourceId": "SRC-MS-02", "categories": ["magnetic-separators"],
        "sourceName": "High Speed Drum Magnets", "organization": "Eriez",
        "sourceType": "manufacturer technical page", "url": "https://www.eriez.com/Products/Magnetic-Separation/Magnetic-Processing-Concentration/Magnetic-Separation-Solutions-for-Dry-Minerals-Processing/DF-High-Speed-Drum-Magnets",
        "evidenceSummary": "Supports the general dry-drum concept that magnetic attraction, gravity, drum speed and material trajectory influence dry separation.",
        "confidence": "medium", "useRestriction": "Generic principle only; no competitor performance, model or imagery is reused.",
    },
    {
        "sourceId": "SRC-MS-03", "categories": ["magnetic-separators"],
        "sourceName": "Suspended Electromagnets", "organization": "Eriez",
        "sourceType": "manufacturer technical page", "url": "https://www.eriez.com/Products/Magnetic-Separation/Magnetic-Separation-for-Tramp-Metal/Large-Magnetic-Separation/Suspended-Electromagnets",
        "evidenceSummary": "Describes conveyor-mounted removal of tramp metal and identifies belt width, speed, burden depth and material type as sizing considerations; self-cleaning configurations use a discharge belt.",
        "confidence": "medium", "useRestriction": "General application framework only; no competitor capacity, efficiency, product or commercial claim is reused.",
    },
    {
        "sourceId": "SRC-MS-04", "categories": ["magnetic-separators"],
        "sourceName": "Eddy Current Separators", "organization": "Eriez",
        "sourceType": "manufacturer technical page", "url": "https://www.eriez.com/Products/Metals-Recycling/Nonferrous-Recovery/Eddy-Current-Separators",
        "evidenceSummary": "Explains that a rotating magnetic rotor induces eddy currents in conductive nonferrous pieces, producing a repelling force that supports physical separation from nonmetals.",
        "confidence": "medium", "useRestriction": "Generic mechanism only; no competitor recovery figures, product data, images or text reused.",
    },
    {
        "sourceId": "SRC-MS-05", "categories": ["magnetic-separators"],
        "sourceName": "Metal Inclusion", "organization": "U.S. Food and Drug Administration",
        "sourceType": "government food-safety guidance", "url": "https://www.fda.gov/downloads/Food/GuidanceRegulation/UCM252440.pdf",
        "evidenceSummary": "Treats a magnet or other separation step as a possible control where metal inclusion is a significant hazard, within a documented hazard-analysis approach.",
        "confidence": "high", "useRestriction": "Does not establish food-grade construction, validation, certification or suitability for any COWIN MACHINE product.",
    },
]


def compact(value: object) -> str:
    return " ".join(str(value or "").split())


def path_only(url: str) -> str:
    return url.replace(SITE, "")


def list_join(values: list[str]) -> str:
    return "; ".join(value for value in values if value) or "Not established"


def read_master():
    master = json.loads((AUDIT_DIR / "canonical-product-master.json").read_text(encoding="utf-8"))
    with (AUDIT_DIR / "product-keywords.csv").open(encoding="utf-8", newline="") as file:
        keywords = {row["canonicalId"]: row for row in csv.DictReader(file)}
    with (AUDIT_DIR / "product-data-quality-issues.csv").open(encoding="utf-8", newline="") as file:
        issues = list(csv.DictReader(file))
    if len(master["products"]) != 135:
        raise ValueError(f"Expected 135 canonical families, found {len(master['products'])}")
    # Explicit reads keep this stage traceable to all three requested audit inputs.
    if not keywords or not issues:
        raise ValueError("Keyword or data-quality source file is empty")
    return master["summary"], master["products"], keywords, issues


def category_terms(category: str) -> dict:
    if category == "compressed-air-equipment":
        return {
            "principle": "A compressor package converts mechanical input into pressurized air. In a rotary screw arrangement, air is compressed through a rotating airend; the downstream system may include cooling, separation, storage, filtration and controls according to the duty.",
            "configuration": ["drive package", "compression element or airend", "cooling and separation stages", "controls and protection", "air receiver and distribution treatment where required"],
            "selection": ["required working pressure at the point of use", "free-air-delivery demand and demand variation", "electric versus diesel power availability", "stationary versus portable installation", "air quality, moisture and filtration requirement", "ambient conditions, access and maintenance plan"],
            "parameters": ["working pressure — pressure available to serve the application, commonly expressed in bar(g) or psi", "free air delivery — delivered air volume normalized to stated reference conditions, commonly m3/min or cfm", "motor or engine power — input power, commonly kW", "duty and control method — load/unload, fixed-speed or variable-speed logic where verified", "air quality / filtration class — treatment requirement selected for the end use"],
            "industries": ["mining and quarrying", "water well drilling", "rock drilling and blasting", "road and infrastructure construction", "manufacturing workshops", "shipyards", "equipment rental"],
            "cases": ["supply air to drilling tools or a DTH system after matching pressure and flow to the rig and hammer", "operate workshop pneumatic tools where electrical supply and air treatment are available", "support temporary site air demand using a portable package after checking transport, ventilation and fuel logistics"],
            "workflow": ["capture point-of-use pressure, flow and air-quality needs", "identify simultaneous demand and peak events", "select power source and portability", "define filtration, condensate and distribution treatment", "commission using measured system pressure and leak checks"],
            "pain": ["pressure loss at distant tools", "insufficient flow during simultaneous demand", "moisture or oil contamination", "fuel or electrical constraints at remote sites", "energy waste from leakage or poor controls"],
            "response": ["A fact-checked selection process aligns pressure and delivered flow with the actual end use and separates equipment selection from air-treatment design.", "System review should include leakage, condensate removal and preventive maintenance rather than only compressor nameplate output."],
            "benefits": ["more appropriate air supply planning", "reduced risk of mismatched pressure or flow", "a clearer basis for maintenance and air-quality discussions"],
            "maintenance": ["follow the approved manufacturer maintenance schedule", "inspect filters, fluid levels and hose connections as applicable", "manage condensate according to the application and local environmental rules", "isolate and depressurize before service", "check leaks and pressure loss as part of routine system review"],
            "related": ["air receiver", "aftercooler", "dryer", "filters", "hoses and couplings", "DTH hammer or pneumatic tools", "condensate treatment equipment"],
            "faqs": ["How should pressure and free-air delivery be matched to a drilling or workshop duty?", "When is diesel power more practical than an electric package?", "What air treatment is needed for the intended process?", "How should a portable package be positioned and ventilated?", "Which model-specific data must be confirmed before purchase?"],
            "sources": ["SRC-CA-01"],
        }
    if category == "generator-systems":
        return {
            "principle": "A diesel generator system combines an engine, alternator, controls and protection to convert fuel energy into electrical power. Selection depends on the intended duty, electrical supply, load profile and installation conditions.",
            "configuration": ["diesel engine", "alternator", "controller and protection", "baseframe or enclosure", "fuel system", "starting system", "output connection and earthing arrangement designed for the site"],
            "selection": ["prime, standby or continuous duty as defined for the project", "running load, starting load and load steps", "phase, voltage and frequency", "ambient temperature and altitude", "open-frame versus enclosed installation", "fuel storage, service access, noise and local compliance requirements"],
            "parameters": ["rated power — commonly expressed in kW and/or kVA under a stated duty rating", "frequency — typically Hz and must match the electrical system", "voltage and phase — output characteristics required by connected loads", "load-step and motor-start response — ability to accommodate transient demand under defined conditions", "fuel consumption — must be quoted only from approved model-specific data and stated load condition"],
            "industries": ["construction sites", "remote infrastructure", "telecom backup", "emergency power", "industrial facilities", "rental fleets", "remote mining operations"],
            "cases": ["support a temporary construction distribution board after load schedule review", "provide remote primary power where utility supply is unavailable", "provide a backup source for selected essential loads with an approved transfer arrangement"],
            "workflow": ["build a load schedule including starting and step loads", "confirm duty classification and operating hours", "confirm voltage, frequency, phase and site conditions", "define installation, fuel and exhaust requirements", "validate the final selection with an electrical engineer for critical duties"],
            "pain": ["undersizing against motor starts", "confusion between prime and standby duty", "voltage or frequency mismatch", "poor fuel and service planning", "noise or enclosure needs not identified early"],
            "response": ["A structured load and duty review keeps the selection focused on the actual power profile rather than a nominal kVA alone.", "Open-frame or enclosed configuration should be assessed against installation, service and environmental needs."],
            "benefits": ["clearer load-sizing inputs", "reduced risk of duty misapplication", "better planning for fuel, access and electrical integration"],
            "maintenance": ["use the approved operation and maintenance manual for the exact set", "inspect fluids, battery, filters, guards and connections before use", "verify earthing, exhaust routing and protective devices with qualified personnel", "test under the required site procedure", "do not assume a rating or load capability without manufacturer confirmation"],
            "related": ["automatic transfer equipment", "distribution board", "cables", "fuel tank", "load bank", "remote monitoring", "acoustic enclosure"],
            "faqs": ["What duty rating matches the operating plan?", "Which load data is required to size the set?", "How are motor starts and load steps considered?", "When is an enclosure required?", "Which voltage and frequency options can be confirmed for this model?"],
            "sources": ["SRC-GEN-01", "SRC-GEN-02"],
        }
    if category == "drilling-equipment":
        return {
            "principle": "A drilling rig applies rotary motion and, depending on the method, percussion, thrust and a flushing medium to create a borehole. The rig, drill string, bit, compressor or mud system must be selected as a matched system for formation conditions and the planned hole.",
            "configuration": ["carrier or crawler undercarriage", "mast and feed system", "rotary head", "power pack", "drill string handling", "air or mud circulation interface", "safety controls and stabilizing equipment"],
            "selection": ["drilling method and formation", "planned hole diameter and depth", "torque, rotation and feed/pull requirements", "air versus mud circulation plan", "terrain, transport and crawler mobility", "compatible hammer, bit, rod and compressor/mud equipment"],
            "parameters": ["hole diameter — borehole size, commonly mm or in", "drilling depth — planned reach, commonly m or ft, subject to formation and tool conditions", "rotary torque — turning capacity, commonly N.m", "rotary speed — rotation rate, rpm", "feed / pull force — thrust and retrieval capability, commonly kN", "air pressure and flow or mud circulation requirement — matched to the drilling system"],
            "industries": ["water well drilling", "agricultural irrigation", "geothermal and ground-source projects", "mining and quarrying", "foundation and anchoring", "geotechnical investigation", "road and railway construction"],
            "cases": ["water-well work that changes method between unconsolidated material and competent bedrock", "DTH drilling in rock after matching hammer, bit, air supply and dust control", "quarry or roadwork drilling where hole layout, geology and downstream blasting requirements are reviewed"],
            "workflow": ["review geology and required borehole outcome", "select circulation method and dust-control plan", "match rig capacity, hammer/bit/rod and compressor or mud system", "prepare stable access and rig-up area", "drill, monitor cuttings and conditions, then inspect consumables and connections"],
            "pain": ["poor tool matching", "hole deviation or unstable formations", "cuttings removal problems", "dust exposure", "transport and terrain constraints", "insufficient verified depth or diameter information"],
            "response": ["Rig selection should treat drilling method, tooling and flushing as one system, with model-specific capacity confirmed from approved documentation.", "The selection discussion should include site conditions and dust-control planning rather than relying on a generic rig label."],
            "benefits": ["more traceable selection inputs", "better tool and air/mud matching", "clearer safety and maintenance planning"],
            "maintenance": ["inspect mast, feed, hoses, guards, wire ropes and structural components before use", "verify pressure ratings and compatible connections", "control dust with an appropriate engineering method for the site", "follow approved lubrication and maintenance procedures", "keep personnel clear of moving drill string and suspended loads"],
            "related": ["air compressor", "mud pump", "DTH hammer", "drill bits", "drill pipes", "casing", "dust collector", "water injection equipment"],
            "faqs": ["Which drilling method matches the formation?", "What information is needed to confirm depth and diameter capability?", "How should a DTH hammer be paired with the bit and air supply?", "What terrain information is needed for crawler selection?", "Which safety controls must be planned before rig-up?"],
            "sources": ["SRC-DR-01", "SRC-DR-02", "SRC-DR-03", "SRC-DR-04"],
        }
    if category == "drilling-consumables":
        return {
            "principle": "Drilling consumables transmit rotary and/or percussion energy to the rock and provide a path for flushing. Their interface, diameter, thread or shank, pressure rating and wear condition must be compatible with the selected rig and drilling method.",
            "configuration": ["drill bit", "hammer or rock drill where relevant", "drill pipe or rod", "shank, thread or sub adapters", "lubrication and flushing interfaces", "handling tools"],
            "selection": ["drilling method", "rig and hammer interface", "hole diameter", "rock formation and abrasiveness", "air pressure and flow where applicable", "thread/shank compatibility", "wear life and safe replacement practice"],
            "parameters": ["bit diameter — cutting size, commonly mm or in", "thread or shank — physical interface that must match the upstream component", "hammer nominal size — tool family reference that must be matched to bit and hole", "working pressure — air-system condition for applicable pneumatic tools", "rod outside diameter and wall thickness — drill-string dimensions, commonly mm", "material / button type — only from approved manufacturer documentation"],
            "industries": ["water well drilling", "mining and quarrying", "foundation and anchoring", "geotechnical investigation", "road and railway construction", "rock drilling"],
            "cases": ["replace worn DTH bits while preserving the required hammer and thread interface", "build a drill string with compatible pipes, adapters and bit", "select pneumatic rock-drill accessories after confirming the exact tool platform"],
            "workflow": ["confirm the rig and drilling method", "verify interface, thread/shank and required hole geometry", "inspect consumables for wear or damage", "install using approved handling tools and torque procedure", "monitor performance and replace before damaged parts create a safety or quality risk"],
            "pain": ["mismatched interfaces", "accelerated wear in abrasive rock", "loss of flushing", "unsafe handling of heavy or pressurized components", "uncertain compatibility caused by generic product names"],
            "response": ["Consumable selection should be based on an exact approved interface and application record; a family name alone is not enough to establish compatibility.", "Wear monitoring and replacement planning reduce the chance of using damaged components."],
            "benefits": ["more controlled tool matching", "clearer replacement planning", "reduced risk of incompatible or worn components"],
            "maintenance": ["inspect threads, shanks, retainers, bits and pipes before assembly", "use manufacturer-approved lubrication and torque practice", "isolate air/hydraulic energy before service", "do not force mismatched components", "remove damaged components from service"],
            "related": ["drilling rig", "air compressor", "DTH hammer", "drill bits", "drill pipes", "lubricator", "adapters", "handling tools"],
            "faqs": ["How is a hammer matched to a bit?", "What thread or shank details are needed for confirmation?", "How does rock condition affect consumable selection?", "When should a bit or rod be removed from service?", "Can a generic model code confirm compatibility?"],
            "sources": ["SRC-DR-02", "SRC-DR-03", "SRC-DR-04"],
        }
    if category == "mobile-lighting-systems":
        return {
            "principle": "A mobile light tower raises luminaires above a work area and supplies electrical energy from a diesel generator, battery, photovoltaic system, hybrid arrangement or a confirmed combination. Actual night operation depends on lighting load, energy storage, solar input, weather, controls and site conditions.",
            "configuration": ["mast and luminaires", "trailer or mobile frame", "stabilizers and jacks", "energy source and power conversion", "battery and charging system where supplied", "controls and monitoring", "CCTV or communications equipment only where approved"],
            "selection": ["required work-area illumination and glare control", "night-time operating duration", "solar availability, shading and seasonal conditions", "diesel, battery, solar or hybrid operating strategy", "mast height, coverage study and wind/site constraints", "trailer transport, levelling, service access and local road rules"],
            "parameters": ["lighting power — electrical demand of luminaires, commonly W", "luminous flux — stated light output, commonly lm", "mast height — raised working height, commonly m or ft", "energy capacity — stored energy, commonly kWh", "runtime — duration under a stated load, state of charge and environment", "solar array rating — PV nameplate capacity, commonly Wp", "wind limit — model-specific operating restriction stated by the manufacturer"],
            "industries": ["construction sites", "roadworks", "mining sites", "remote camps", "emergency response", "event lighting", "security and CCTV monitoring", "rental fleets"],
            "cases": ["temporary night construction lighting after a site lighting plan and glare review", "remote-area lighting where solar yield, battery storage and backup power are assessed", "security or CCTV monitoring only after confirming camera, communications and power configuration"],
            "workflow": ["define required task illumination and operating hours", "review location, shadows, weather and solar access", "choose energy strategy and verify storage/runtime assumptions", "transport, level and stabilize the unit as instructed for the exact model", "aim luminaires, test controls and monitor operating conditions"],
            "pain": ["runtime assumptions not tied to actual load", "insufficient solar charging conditions", "glare or uneven coverage", "wind, slope or overhead-line risk", "unverified CCTV or telematics configuration"],
            "response": ["A selection brief should separate illumination need from energy-source assumptions and document the actual night-time load profile.", "Solar, battery, diesel and hybrid options should be compared against site energy availability, transport and service constraints."],
            "benefits": ["more realistic runtime planning", "clearer energy-source comparison", "better preparation for safe placement and lighting design"],
            "maintenance": ["use the exact approved operating manual for mast, stabilizer and trailer procedures", "inspect electrical cables, guards, hitch and stabilizers before use", "keep clear of overhead power lines and do not exceed model-specific wind limits", "manage battery systems under the approved charging and safety procedures", "follow local illumination and work-zone requirements"],
            "related": ["mobile generator", "battery storage", "PV modules", "trailer equipment", "remote monitoring", "CCTV package", "distribution cables"],
            "faqs": ["How should solar, diesel, battery and hybrid options be compared?", "What load assumptions are required to calculate runtime?", "How is mast height selected for a lighting plan?", "What site checks are needed before raising the mast?", "Which CCTV and monitoring functions are confirmed for this exact product?"],
            "sources": ["SRC-LT-01", "SRC-LT-02", "SRC-LT-03", "SRC-LT-04"],
        }
    return {
        "principle": "Magnetic separation uses a magnetic field or, for eddy-current separation, an induced electromagnetic effect, to separate targeted materials from a product stream. The equipment type must be selected for material properties, feed condition and installation geometry.",
        "configuration": ["magnetic circuit or electromagnet", "support frame or housing", "conveyor, drum, chute or process vessel interface", "manual or self-cleaning discharge system where applicable", "drive and controls where applicable", "feed-conditioning equipment"],
        "selection": ["ferrous versus nonferrous recovery or tramp-metal removal objective", "permanent versus electromagnetic field requirement", "manual-clean versus self-cleaning duty", "belt width, burden depth, suspension height and speed for conveyor installations", "dry versus wet feed condition", "particle size, liberation, capacity and material-flow behavior", "food/process hygiene requirements only where independently verified"],
        "parameters": ["magnetic field intensity — must state measurement method and location, commonly mT or gauss", "belt or working width — process cross-section, commonly mm", "suspension height — distance between magnet face and conveyor/burden reference, commonly mm", "burden depth — material layer depth, commonly mm", "capacity — throughput under stated material conditions, commonly t/h", "drum diameter / speed — mechanical separation geometry, commonly mm and rpm", "slurry solids and feed size — wet-process conditions that affect separation"],
        "industries": ["mining and mineral processing", "aggregate and quarrying", "cement", "coal handling", "recycling and waste sorting", "scrap metal processing", "food and grain processing", "plastics and chemicals", "power plants and bulk material handling"],
        "cases": ["remove tramp ferrous metal ahead of crushers, mills or shredders after verifying conveyor geometry and burden conditions", "treat dry mineral feed where particle size, feed presentation and magnetic response are established", "treat slurry with a wet-drum process after laboratory/process confirmation", "recover conductive nonferrous material by eddy-current separation after upstream ferrous removal and feed preparation"],
        "workflow": ["characterize feed material and separation target", "select dry, wet, suspended, drum, pulley, filter or eddy-current approach", "record installation geometry and material-flow conditions", "test or validate separation performance with representative material", "inspect collection/discharge and maintain safe cleaning procedures"],
        "pain": ["tramp metal damaging downstream equipment", "poor separation due to deep or uneven burden", "wrong dry/wet configuration", "manual cleaning workload", "unverified food-grade or process-hygiene claims", "unknown feed behavior or material response"],
        "response": ["Selection should begin with the separation objective and feed characterization, then confirm installation geometry and duty before choosing a magnetic system.", "Manual versus self-cleaning operation should be tied to contamination frequency, access and safe discharge planning."],
        "benefits": ["a traceable basis for selecting a separator type", "better protection planning for downstream equipment", "clearer separation-test and installation inputs"],
        "maintenance": ["isolate power and follow approved lockout procedures for electromagnets and drives", "keep hands and ferromagnetic tools clear of strong magnetic fields", "inspect belts, scrapers, guards, bearings and supports as applicable", "use safe collection and discharge procedures for captured metal", "do not claim hygienic or food-grade suitability without verified construction and validation evidence"],
        "related": ["metal detector", "vibratory feeder", "conveyor", "crusher or mill protection", "drum separator", "eddy-current separator", "magnetic filter or grid", "collection bin"],
        "faqs": ["When should a permanent or electromagnetic separator be considered?", "When is self-cleaning preferred over manual cleaning?", "Which conveyor dimensions and burden information are needed?", "How is dry separation different from wet slurry separation?", "What evidence is required before food-processing suitability is stated?"],
        "sources": ["SRC-MS-01", "SRC-MS-02", "SRC-MS-03", "SRC-MS-04", "SRC-MS-05"],
    }


def source_ids_for(product: dict, profile: dict) -> list[str]:
    family = (product["family"] + " " + (product.get("model") or "")).lower()
    ids = list(profile["sources"])
    if product["category"] == "magnetic-separators":
        ids = ["SRC-MS-05"] if "food" in family or "grain" in family else []
        if "wet" in family:
            ids.append("SRC-MS-01")
        elif "dry" in family or "drum" in family:
            ids.append("SRC-MS-02")
        elif "eddy" in family:
            ids.append("SRC-MS-04")
        elif any(word in family for word in ["suspended", "overband", "rcyd", "rcdd", "rcyb", "rcdb", "rcda", "rcdc", "rcde", "rcdf"]):
            ids.append("SRC-MS-03")
        else:
            ids.extend(["SRC-MS-02", "SRC-MS-03"])
        return list(dict.fromkeys(ids))
    if product["category"] == "drilling-consumables" and "dth" not in family and "hammer" not in family:
        return ["SRC-DR-04"]
    return ids


def alternative_terms(product: dict, keyword_row: dict | None) -> list[str]:
    family = product["family"]
    model = product.get("model")
    category = product["category"]
    terms = [family]
    if model:
        terms.extend([model, f"{model} {family}"])
    additions = {
        "compressed-air-equipment": ["industrial air compressor", "compressed air equipment", "pressure and free-air-delivery selection"],
        "generator-systems": ["diesel generator", "generator set", "temporary power generator", "prime power vs standby power"],
        "drilling-equipment": ["drilling rig", "water well drilling rig", "crawler drilling rig", "DTH drilling equipment"],
        "drilling-consumables": ["drilling tool", "drilling consumable", "DTH hammer and bit matching", "drill pipe compatibility"],
        "mobile-lighting-systems": ["mobile light tower", "temporary site lighting", "solar diesel battery hybrid light tower"],
        "magnetic-separators": ["magnetic separation equipment", "tramp iron removal", "material separation equipment"],
    }
    keyword_terms = []
    if keyword_row:
        for key in ("primary", "secondary", "longTail", "buyerIntent"):
            keyword_terms.extend([value.strip() for value in keyword_row.get(key, "").split("|") if value.strip()])
    return list(dict.fromkeys(terms + additions[category] + keyword_terms))


def status_mode(product: dict) -> tuple[str, str]:
    if product["productStatus"] == "verified-model":
        return ("model-identified research profile", "moderate")
    return ("requires-owner-confirmation card", "low")


def high_risk_items(product: dict) -> list[str]:
    base = [
        "Any exact capacity, output, range, consumption, recovery, runtime, coverage, noise, pressure, size, weight or operating limit not supported by an approved model-specific manufacturer source.",
        "Any certification, food-grade, explosion-proof, environmental, compliance, warranty, lead-time, country-of-origin, availability or after-sales statement not approved by the owner.",
        "Any claim that this model is self-manufactured, exclusive, compatible with another supplier’s tool, or suitable for a particular site without documented confirmation.",
    ]
    if product["productStatus"] != "verified-model":
        base.insert(0, "The model or family identity itself is unresolved; do not publish technical product-page content beyond a neutral owner-confirmation placeholder.")
    return base


def current_source_entries(product: dict) -> list[dict]:
    entries = []
    for index, url in enumerate(product["currentUrls"], 1):
        entries.append({
            "sourceId": f"SRC-CURRENT-{product['canonicalId']}-{index}",
            "canonicalId": product["canonicalId"],
            "category": product["category"],
            "sourceName": "COWIN MACHINE public product record",
            "organization": "COWIN MACHINE public website",
            "sourceType": "current-site audit evidence",
            "url": url,
            "accessedAt": ACCESS_DATE,
            "evidenceSummary": "Current URL, H1/model reference and visible specification table captured in the preceding product audit. This is traceability evidence, not independent manufacturer verification.",
            "confidence": "low",
            "useRestriction": "Do not treat as authoritative model documentation until the owner supplies approved manufacturer evidence.",
        })
    return entries


def build_dossier(product: dict, source_lookup: dict, keyword_row: dict | None, issue_rows: list[dict]) -> dict:
    profile = category_terms(product["category"])
    mode, confidence = status_mode(product)
    external_ids = source_ids_for(product, profile)
    current_sources = current_source_entries(product)
    current_ids = [entry["sourceId"] for entry in current_sources]
    all_source_ids = current_ids + external_ids
    issue_statuses = sorted({row["status"] for row in issue_rows if row.get("url") in product["currentUrls"]})
    missing = list(product.get("missingSpecifications") or [])
    if product["productStatus"] != "verified-model":
        missing = list(dict.fromkeys(missing + ["Owner-approved model/family identity", "Model-specific manufacturer documentation", "Authorized product-image rights and current specification evidence"]))
    captured_specs = product.get("verifiedSpecifications") or {}
    publishing_status = "requires-owner-confirmation" if captured_specs else "not-available"
    dossier = {
        "schemaVersion": "1.0",
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "dossierMode": mode,
        "productIdentity": {
            "canonicalId": product["canonicalId"], "category": product["category"], "categoryLabel": CATEGORY_LABELS[product["category"]],
            "family": product["family"], "model": product.get("model"), "variant": product.get("variant"),
            "currentUrls": product["currentUrls"], "aliases": product["aliases"], "auditStatus": product["productStatus"],
        },
        "alternativeNamesAndSearchTerms": alternative_terms(product, keyword_row),
        "auditDataQualityIssues": issue_statuses,
        "workingPrinciple": {"content": profile["principle"], "evidenceSourceIds": external_ids, "scope": "category-general; not model-specific"},
        "typicalConfiguration": profile["configuration"],
        "mainComponents": profile["configuration"],
        "buyerSelectionFactors": profile["selection"],
        "commonTechnicalParameters": profile["parameters"],
        "parameterDefinitionsAndUnits": profile["parameters"],
        "typicalOperatingConditions": [
            "Use only within the exact model’s approved environmental, duty and installation limits.",
            "Confirm site utilities, material/formation conditions, access, local safety requirements and maintenance access before selection.",
            "Do not convert category-general research into a model-specific operating envelope without approved documentation.",
        ],
        "suitableIndustries": profile["industries"],
        "detailedUseCases": profile["cases"],
        "operationalWorkflow": profile["workflow"],
        "buyerPainPoints": profile["pain"],
        "howProductAddressesPainPoints": profile["response"],
        "expectedOperationalBenefits": {"content": profile["benefits"], "qualification": "Planning benefits only; no quantified, model-specific or site-specific performance outcome is claimed."},
        "maintenanceAndSafetyConsiderations": profile["maintenance"],
        "relatedEquipmentAndConsumables": profile["related"],
        "commonBuyerFaqs": profile["faqs"],
        "currentWebsiteCapturedSpecifications": {
            "values": captured_specs,
            "publicationStatus": publishing_status,
            "note": "Captured from the current public product record during audit. These values are not independently verified manufacturer data and must be owner-confirmed before technical publication.",
            "evidenceSourceIds": current_ids,
        },
        "missingOrRequiresOwnerConfirmation": missing,
        "unpublishableHighRiskContent": high_risk_items(product),
        "originalDetailPageContentFramework": [
            "Original positioning paragraph based on confirmed application and product identity.",
            "Working-principle explanation with a source-backed, non-brand-specific diagram brief.",
            "Selection factors and parameter definitions; publish exact values only from approved fact cards.",
            "Application workflow, maintenance and safety content adapted to local requirements and the approved manual.",
            "Buyer FAQ that answers procurement questions without copying supplier wording or promising unverified outcomes.",
        ],
        "evidenceSourceList": [
            {"sourceId": source_id, "sourceName": source_lookup[source_id]["sourceName"], "url": source_lookup[source_id]["url"], "accessedAt": source_lookup[source_id]["accessedAt"], "evidenceSummary": source_lookup[source_id]["evidenceSummary"], "confidence": source_lookup[source_id]["confidence"]}
            for source_id in all_source_ids
        ],
        "contentConfidenceLevel": {
            "overall": confidence,
            "modelSpecific": "requires-owner-confirmation",
            "explanation": "External sources support only category-level principles and selection logic. The current website record preserves traceability but does not replace approved manufacturer documentation for a model-specific claim.",
        },
    }
    return dossier


def dossier_markdown(dossier: dict) -> str:
    identity = dossier["productIdentity"]
    def bullets(values): return "\n".join(f"- {value}" for value in values)
    sources = "\n".join(f"- `{source['sourceId']}` — {source['sourceName']} ({source['accessedAt']}): {source['url']}\n  Evidence: {source['evidenceSummary']}\n  Confidence: {source['confidence']}" for source in dossier["evidenceSourceList"])
    params = bullets(dossier["commonTechnicalParameters"])
    captured = "\n".join(f"- {key}: {value}" for key, value in dossier["currentWebsiteCapturedSpecifications"]["values"].items()) or "- No captured values available."
    return f"""# {identity['family']}{' — ' + identity['model'] if identity['model'] else ''}

**Canonical ID:** {identity['canonicalId']}  
**Category:** {identity['categoryLabel']}  
**Research mode:** {dossier['dossierMode']}  
**Audit status:** {identity['auditStatus']}

## 1. Product identity

Current URL mapping: {', '.join(identity['currentUrls'])}

## 2. Alternative names and search terms

{bullets(dossier['alternativeNamesAndSearchTerms'])}

## 3. Working principle

{dossier['workingPrinciple']['content']}

## 4. Typical configuration

{bullets(dossier['typicalConfiguration'])}

## 5. Main components

{bullets(dossier['mainComponents'])}

## 6. Buyer selection factors

{bullets(dossier['buyerSelectionFactors'])}

## 7–8. Common technical parameters, definitions and units

{params}

## 9. Typical operating conditions

{bullets(dossier['typicalOperatingConditions'])}

## 10–11. Suitable industries and detailed use cases

**Industries**

{bullets(dossier['suitableIndustries'])}

**Use cases**

{bullets(dossier['detailedUseCases'])}

## 12. Operational workflow

{bullets(dossier['operationalWorkflow'])}

## 13–15. Buyer pain points, response and expected benefits

**Pain points**

{bullets(dossier['buyerPainPoints'])}

**Selection response**

{bullets(dossier['howProductAddressesPainPoints'])}

**Expected operational benefits**

{bullets(dossier['expectedOperationalBenefits']['content'])}

{dossier['expectedOperationalBenefits']['qualification']}

## 16. Maintenance and safety considerations

{bullets(dossier['maintenanceAndSafetyConsiderations'])}

## 17. Related equipment and consumables

{bullets(dossier['relatedEquipmentAndConsumables'])}

## 18. Common buyer FAQs

{bullets(dossier['commonBuyerFaqs'])}

## 19. Evidence source list

{sources}

## 20. Content confidence level

Overall: **{dossier['contentConfidenceLevel']['overall']}**.  
Model-specific: **{dossier['contentConfidenceLevel']['modelSpecific']}**.

{dossier['contentConfidenceLevel']['explanation']}

## Captured current-site specification values — not publication-ready

{captured}

{dossier['currentWebsiteCapturedSpecifications']['note']}

## Unpublishable / owner-confirmation items

{bullets(dossier['missingOrRequiresOwnerConfirmation'] + dossier['unpublishableHighRiskContent'])}

## Original product-detail content framework

{bullets(dossier['originalDetailPageContentFramework'])}
"""


def write_data(summary, products, keyword_lookup, all_issues):
    RESEARCH_DIR.mkdir(parents=True, exist_ok=True)
    DOSSIER_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_RESEARCH_DIR.mkdir(parents=True, exist_ok=True)
    source_lookup = {source["sourceId"]: {**source, "accessedAt": ACCESS_DATE} for source in EXTERNAL_SOURCES}
    all_current_entries = []
    dossiers = []
    for product in products:
        entries = current_source_entries(product)
        all_current_entries.extend(entries)
        for entry in entries:
            source_lookup[entry["sourceId"]] = entry
        relevant_issues = [issue for issue in all_issues if issue.get("url") in product["currentUrls"]]
        dossier = build_dossier(product, source_lookup, keyword_lookup.get(product["canonicalId"]), relevant_issues)
        dossiers.append(dossier)
        (DOSSIER_DIR / f"{product['canonicalId']}.json").write_text(json.dumps(dossier, indent=2, ensure_ascii=False), encoding="utf-8")
        (DOCS_RESEARCH_DIR / f"{product['canonicalId']}.md").write_text(dossier_markdown(dossier), encoding="utf-8")

    external_rows = [{
        **{key: value for key, value in source.items() if key != "categories"},
        "category": " | ".join(source["categories"]),
        "accessedAt": ACCESS_DATE,
        "canonicalId": "category-general",
    } for source in EXTERNAL_SOURCES]
    source_rows = external_rows + all_current_entries
    headers = ["sourceId", "canonicalId", "category", "sourceName", "organization", "sourceType", "url", "accessedAt", "evidenceSummary", "confidence", "useRestriction"]
    with (RESEARCH_DIR / "source-register.csv").open("w", encoding="utf-8", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=headers)
        writer.writeheader()
        writer.writerows(source_rows)

    matrix = {"generatedAt": datetime.now(timezone.utc).isoformat(), "scope": "Research-only; no online website changes", "categories": []}
    keyword_rows = []
    for category in CATEGORY_LABELS:
        category_dossiers = [dossier for dossier in dossiers if dossier["productIdentity"]["category"] == category]
        profile = category_terms(category)
        matrix["categories"].append({
            "category": category, "label": CATEGORY_LABELS[category], "industries": profile["industries"], "useCases": profile["cases"], "selectionFactors": profile["selection"], "operationalWorkflow": profile["workflow"], "safetyConsiderations": profile["maintenance"],
            "sourceIds": list(dict.fromkeys(category_terms(category)["sources"] + [source_id for dossier in category_dossiers for source_id in [source["sourceId"] for source in dossier["evidenceSourceList"]] if not source_id.startswith("SRC-CURRENT-")])),
            "productFamilies": [{"canonicalId": dossier["productIdentity"]["canonicalId"], "family": dossier["productIdentity"]["family"], "model": dossier["productIdentity"]["model"], "auditStatus": dossier["productIdentity"]["auditStatus"], "contentConfidence": dossier["contentConfidenceLevel"]["overall"]} for dossier in category_dossiers],
        })
        for dossier in category_dossiers:
            identity = dossier["productIdentity"]
            keyword_rows.append({
                "canonicalId": identity["canonicalId"], "category": category, "productFamily": identity["family"], "model": identity["model"] or "", "auditStatus": identity["auditStatus"],
                "industries": " | ".join(dossier["suitableIndustries"]), "applicationSearchTerms": " | ".join(dossier["alternativeNamesAndSearchTerms"]), "buyerSelectionFactors": " | ".join(dossier["buyerSelectionFactors"]), "publicationStatus": dossier["contentConfidenceLevel"]["modelSpecific"],
            })
    (RESEARCH_DIR / "industry-use-case-matrix.json").write_text(json.dumps(matrix, indent=2, ensure_ascii=False), encoding="utf-8")
    with (RESEARCH_DIR / "industry-keyword-map.csv").open("w", encoding="utf-8", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=list(keyword_rows[0]))
        writer.writeheader(); writer.writerows(keyword_rows)
    return dossiers, source_rows, matrix


# Compact Word helpers using the compact_reference_guide design preset.
def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"; run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def setup_doc(doc):
    for section in doc.sections:
        section.top_margin = Inches(1); section.bottom_margin = Inches(1); section.left_margin = Inches(1); section.right_margin = Inches(1); section.header_distance = Inches(.492); section.footer_distance = Inches(.492)
        section.header.is_linked_to_previous = False; section.footer.is_linked_to_previous = False
        h = section.header.paragraphs[0]; h.text = ""; set_font(h.add_run("COWIN MACHINE  |  Industry and Application Research"), 8.5, MUTED, True)
        f = section.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.RIGHT; f.text = ""; set_font(f.add_run("Research-only library  |  Page "), 8.5, MUTED)
        r = f.add_run(); begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin"); instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"; end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); r._r.extend([begin, instr, end])
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri"); normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK,10,5)]:
        style = doc.styles[name]; style.font.name = "Calibri"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color); style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True


def paragraph(doc, value, style=None, fill=None, bold_prefix=None, size=None, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    if fill:
        ppr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); shd.set(qn("w:val"), "clear"); ppr.append(shd)
    if bold_prefix and value.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), size, color, True); set_font(p.add_run(value[len(bold_prefix):]), size, color)
    else: set_font(p.add_run(value), size, color)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.2; set_font(p.add_run(item))


def shade(cell, fill):
    tc = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); shd.set(qn("w:val"), "clear"); tc.append(shd)


def table(doc, headers, rows, widths, font_size=8):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; header = t.rows[0]
    trpr = header._tr.get_or_add_trPr(); marker = OxmlElement("w:tblHeader"); marker.set(qn("w:val"), "true"); trpr.append(marker)
    for i, value in enumerate(headers):
        c = header.cells[i]; c.text = ""; c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; shade(c, HEADER); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0); set_font(p.add_run(value), font_size, INK, True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]; c.text = ""; c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0; set_font(p.add_run(compact(value)), font_size, INK)
    apply_table_geometry(t, widths, table_width_dxa=sum(widths), indent_dxa=120, cell_margins_dxa={"top":80,"bottom":80,"start":120,"end":120})
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return t


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}"); set_font(p.add_run(text), {1:16,2:13,3:12}[level], {1:BLUE,2:BLUE,3:DARK}[level], True); return p


def build_report(summary, dossiers, source_rows, matrix):
    doc = Document(); setup_doc(doc)
    paragraph(doc, "GLOBAL TECHNICAL RESEARCH", size=11, color="7A5A00", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "COWIN MACHINE", size=30, color=INK, bold_prefix="COWIN MACHINE", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Industry and Application Research", size=18, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "Research-only product-family dossiers, technical principles and publication guardrails", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(52)
    paragraph(doc, f"Generated: {ACCESS_DATE}", size=10, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "No website pages, product data, assets, source code or deployment configuration were modified.", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()
    status_counts = Counter(d["productIdentity"]["auditStatus"] for d in dossiers)
    heading(doc, "1. Research scope and evidence rules")
    paragraph(doc, f"This local research library expands the prior audit of {summary['totalRawRecords']} URL records into {len(dossiers)} canonical product-family dossiers. It distinguishes category-level technical research from model-specific claims. External technical sources support operating principles, selection factors, safety and application context; they do not establish COWIN MACHINE product specifications, identity, performance, certification or suitability.")
    bullets(doc, ["No competitor logo, image, PDF, certificate, case study, company history, customer claim, pricing, warranty or marketing copy is reused.", "Every captured current-site specification is explicitly retained as requires-owner-confirmation rather than publication-ready evidence.", "Generic-series, duplicate, possible-misclassification and specification-gap families receive a confirmation card, not a technical product promise.", "Final product-page content must be original and only use numbers or technical commitments from an approved model-specific fact card."])
    table(doc, ["Audit status", "Families", "Research treatment"], [[status, str(count), "Model-identified research profile" if status == "verified-model" else "Owner-confirmation card; no model-specific commitment"] for status, count in sorted(status_counts.items())], [2500,1100,5760], 8.5)
    heading(doc, "2. Industry and use-case matrix")
    matrix_rows = []
    for cat in matrix["categories"]:
        matrix_rows.append([cat["label"], "; ".join(cat["industries"]), "; ".join(cat["useCases"]), "; ".join(cat["sourceIds"])])
    table(doc, ["Product category", "Industries", "Representative use cases", "External source IDs"], matrix_rows, [1750,2850,3250,1510], 7.4)
    heading(doc, "3. External technical source register")
    paragraph(doc, "The following sources were opened and reviewed. They are used as external technical context only; no protected material or commercial statement is copied into COWIN MACHINE content.")
    external = [row for row in source_rows if not row["sourceId"].startswith("SRC-CURRENT-")]
    table(doc, ["ID", "Source / organization", "Category", "Evidence use", "Confidence"], [[row["sourceId"], f"{row['sourceName']} — {row['organization']}", "; ".join(CATEGORY_LABELS.get(c, c) for c in row["category"].split(" | ")), row["evidenceSummary"], row["confidence"]] for row in external], [850,2200,1600,3850,860], 7.2)
    heading(doc, "4. Product-family technical dossiers")
    paragraph(doc, "Each entry below is an original research framework. Model-specific details remain constrained by the current audit status and the evidence guardrails stated with each entry.", fill=NOTE)
    for category in CATEGORY_LABELS:
        category_dossiers = [d for d in dossiers if d["productIdentity"]["category"] == category]
        heading(doc, CATEGORY_LABELS[category], 2)
        for dossier in category_dossiers:
            ident = dossier["productIdentity"]
            title = f"{ident['family']}{' — ' + ident['model'] if ident['model'] else ''} ({ident['canonicalId']})"
            heading(doc, title, 3)
            paragraph(doc, f"Research mode: {dossier['dossierMode']} | Audit status: {ident['auditStatus']} | Confidence: {dossier['contentConfidenceLevel']['overall']}", size=8.8, color=MUTED)
            paragraph(doc, dossier["workingPrinciple"]["content"])
            paragraph(doc, "Selection logic: " + "; ".join(dossier["buyerSelectionFactors"][:5]), bold_prefix="Selection logic: ")
            paragraph(doc, "Typical parameters and units: " + "; ".join(dossier["commonTechnicalParameters"][:4]), bold_prefix="Typical parameters and units: ")
            paragraph(doc, "Applicable industries and scenarios: " + "; ".join(dossier["suitableIndustries"][:5]) + ". " + " ".join(dossier["detailedUseCases"][:2]), bold_prefix="Applicable industries and scenarios: ")
            paragraph(doc, "Maintenance and safety: " + "; ".join(dossier["maintenanceAndSafetyConsiderations"][:4]), bold_prefix="Maintenance and safety: ")
            paragraph(doc, "Original detail-page framework: " + "; ".join(dossier["originalDetailPageContentFramework"][:3]), bold_prefix="Original detail-page framework: ")
            paragraph(doc, "Owner confirmation required: " + "; ".join(dossier["missingOrRequiresOwnerConfirmation"][:3]), fill=RISK, bold_prefix="Owner confirmation required: ")
            paragraph(doc, "Evidence: " + "; ".join(f"{source['sourceId']} ({source['confidence']})" for source in dossier["evidenceSourceList"]), size=8.4, color=MUTED, bold_prefix="Evidence: ")
    heading(doc, "5. Content that must not be published without owner confirmation")
    blocked = [d for d in dossiers if d["productIdentity"]["auditStatus"] != "verified-model"]
    paragraph(doc, f"{len(blocked)} of {len(dossiers)} canonical families are not eligible for model-specific technical claims. Their identity or consolidation status, model evidence, or specification evidence requires review.")
    table(doc, ["Canonical ID", "Family / model", "Audit status", "Publication control"], [[d["productIdentity"]["canonicalId"], d["productIdentity"]["family"] + (f" — {d['productIdentity']['model']}" if d["productIdentity"]["model"] else ""), d["productIdentity"]["auditStatus"], "No model-specific parameters, certification, performance, compatibility or suitability claims"] for d in blocked], [2000,3300,1700,2360], 7.4)
    heading(doc, "6. Next research and content steps")
    bullets(doc, ["Ask the owner to approve the canonical identity and duplicate consolidation decisions before changing any public URL or product content.", "Obtain manufacturer-approved fact cards containing model, exact specification values, applicable conditions, authorized imagery and document rights.", "Use laboratory/process testing or written application engineering confirmation before magnetic separation, lighting coverage, drilling depth, generator loading or compressor delivery claims are published.", "Create original page briefs from the individual dossier frameworks, then run a separate source-to-claim review before implementation."])
    doc.core_properties.title = "COWIN MACHINE Industry and Application Research"
    doc.core_properties.subject = "Research-only technical dossier library"
    doc.core_properties.author = "COWIN MACHINE research workflow"
    doc.save(REPORT_PATH)


def main():
    summary, products, keyword_lookup, all_issues = read_master()
    dossiers, source_rows, matrix = write_data(summary, products, keyword_lookup, all_issues)
    build_report(summary, dossiers, source_rows, matrix)
    status_by_category = {}
    for category in CATEGORY_LABELS:
        items = [d for d in dossiers if d["productIdentity"]["category"] == category]
        status_by_category[category] = {
            "verifiedProductFamilies": sum(d["productIdentity"]["auditStatus"] == "verified-model" for d in items),
            "requiresOwnerConfirmation": sum(d["productIdentity"]["auditStatus"] != "verified-model" for d in items),
            "externalSources": len(set(category_terms(category)["sources"]) | set(source["sourceId"] for d in items for source in d["evidenceSourceList"] if not source["sourceId"].startswith("SRC-CURRENT-"))),
            "detailPageContentFrameworks": sum(d["productIdentity"]["auditStatus"] == "verified-model" for d in items),
            "highRiskUnpublishableFamilies": sum(d["productIdentity"]["auditStatus"] != "verified-model" for d in items),
        }
    print(json.dumps({"dossiers": len(dossiers), "sourceRegisterRows": len(source_rows), "report": str(REPORT_PATH), "categorySummary": status_by_category}, indent=2))


if __name__ == "__main__":
    main()
