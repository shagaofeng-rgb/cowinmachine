"""Create the local COWIN MACHINE product research report from audit JSON.

This script intentionally writes only research deliverables under data/ and docs/.
It does not touch application routes, public assets, deployments, or product source data.
"""

from __future__ import annotations

import csv
import json
import sys
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
AUDIT_DIR = ROOT / "data" / "product-audit"
DOCS_DIR = ROOT / "docs"
DOCX_PATH = DOCS_DIR / "COWIN-MACHINE-Product-Research-Report.docx"
MD_PATH = DOCS_DIR / "COWIN-MACHINE-Product-Research-Report.md"
SKILL_SCRIPTS = Path(r"D:\codex\.codex\plugins\cache\openai-primary-runtime\documents\26.812.11052\skills\documents\scripts")
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WARNING_FILL = "FFF4E5"
WHITE = "FFFFFF"
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

CATEGORY_LABELS = {
    "compressed-air-equipment": "Compressed Air Equipment",
    "generator-systems": "Generator Systems",
    "drilling-equipment": "Drilling Equipment",
    "drilling-consumables": "Drilling Tools & Consumables",
    "mobile-lighting-systems": "Solar & Mobile Light Towers",
    "magnetic-separators": "Magnetic Separators",
}

RECOMMENDED_FIELDS = {
    "compressed-air-equipment": "Approved model; drive and power source; working pressure; free air delivery; airend; dimensions; net weight; applications; authorized images; manual/datasheet reference.",
    "generator-systems": "Approved model; rated/prime/standby power where verified; frequency; voltage; engine and alternator references; enclosure type; dimensions; net weight; applications; authorized images.",
    "drilling-equipment": "Approved model; drilling method; hole diameter range; depth range; power source; travel/chassis configuration; dimensions; net weight; compatible tooling; authorized images.",
    "drilling-consumables": "Approved family/model; tool interface; diameter and thread or shank; compatible rig/hammer; material or button type where verified; packing details where verified; authorized images.",
    "mobile-lighting-systems": "Approved model; energy source; light type and quantity; mast height; battery/storage data where verified; runtime/testing condition; trailer configuration; CCTV option where verified; dimensions; authorized images.",
    "magnetic-separators": "Approved model; separation method; magnetic system; belt/drum/working width; capacity or material condition where verified; installation configuration; power where relevant; dimensions; authorized images.",
}


def text(value: object) -> str:
    return " ".join(str(value or "").split())


def short(value: object, limit: int = 150) -> str:
    value = text(value)
    return value if len(value) <= limit else f"{value[: limit - 1].rstrip()}…"


def path_only(url: str) -> str:
    return url.replace("https://cowinmachine.com", "")


def join(values: list[str], sep: str = "; ") -> str:
    return sep.join(text(value) for value in values if text(value)) or "—"


def read_data():
    raw_doc = json.loads((AUDIT_DIR / "raw-product-inventory.json").read_text(encoding="utf-8"))
    canonical_doc = json.loads((AUDIT_DIR / "canonical-product-master.json").read_text(encoding="utf-8"))
    with (AUDIT_DIR / "product-data-quality-issues.csv").open(encoding="utf-8", newline="") as file:
        issues = list(csv.DictReader(file))
    return raw_doc["summary"], raw_doc["records"], canonical_doc["products"], issues


def _set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_page_number(paragraph):
    run = paragraph.add_run()
    _set_run_font(run, 8.5, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def set_page(section, landscape: bool = False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def build_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Report Small" not in doc.styles:
        small = doc.styles.add_style("Report Small", WD_STYLE_TYPE.PARAGRAPH)
        small.base_style = doc.styles["Normal"]
        small.font.size = Pt(8.5)
        small.paragraph_format.space_after = Pt(3)
        small.paragraph_format.line_spacing = 1.05
    if "Report Caption" not in doc.styles:
        caption = doc.styles.add_style("Report Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption.base_style = doc.styles["Normal"]
        caption.font.size = Pt(8.5)
        caption.font.color.rgb = RGBColor.from_string(MUTED)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(4)
        caption.paragraph_format.line_spacing = 1.0


def add_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("COWIN MACHINE  |  Product Research Report")
    _set_run_font(run, 8.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Research-only local audit  |  Page ")
    _set_run_font(run, 8.5, MUTED)
    add_page_number(footer)


def add_para(doc: Document, value: str, *, style: str | None = None, bold_prefix: str | None = None, fill: str | None = None, align=None, size: float | None = None, color: str | None = None, italic: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if fill:
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        shading.set(qn("w:val"), "clear")
        p_pr.append(shading)
    if bold_prefix and value.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        _set_run_font(first, size=size, color=color, bold=True)
        rest = paragraph.add_run(value[len(bold_prefix):])
        _set_run_font(rest, size=size, color=color, italic=italic)
    else:
        run = paragraph.add_run(value)
        _set_run_font(run, size=size, color=color, italic=italic)
    return paragraph


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        _set_run_font(paragraph.add_run(item))


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        _set_run_font(paragraph.add_run(item))


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], *, font_size: float = 8, header_fill: str = HEADER_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header = table.rows[0]
    set_repeat_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        _set_run_font(run, font_size, WHITE if header_fill == DARK_BLUE else INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if index in (0, len(row) - 1) and len(value) < 40:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            _set_run_font(run, font_size, INK)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=TABLE_INDENT, cell_margins_dxa=CELL_MARGINS)
    doc.add_paragraph(style="Report Caption")
    return table


def add_section_heading(doc: Document, number: int, title: str):
    paragraph = doc.add_paragraph(style="Heading 1")
    _set_run_font(paragraph.add_run(f"{number}. {title}"), 16, BLUE, bold=True)


def inventory_rows(records):
    result = []
    for index, record in enumerate(records, 1):
        details = f"Specs: {len(record['visibleSpecifications'])}; image: {'yes' if record['imagePath'] else 'no'}; related: {len(record['relatedLinks'])}; lastmod: {record['lastModified'] or 'not available'}"
        result.append([
            str(index),
            CATEGORY_LABELS.get(record["category"], record["category"]),
            join([record.get("currentH1") or "H1 not captured", f"Model: {record.get('modelReference') or 'not captured'}"], "\n"),
            details,
            join(record.get("statuses") or []),
            path_only(record["url"]),
        ])
    return result


def canonical_rows(products):
    result = []
    for product in products:
        result.append([
            product["canonicalId"],
            CATEGORY_LABELS.get(product["category"], product["category"]),
            product["family"],
            product.get("model") or "—",
            str(len(product["rawRecordIds"])),
            product["productStatus"],
            join([path_only(url) for url in product["currentUrls"]], "\n"),
        ])
    return result


def keyword_rows(products):
    result = []
    for product in products:
        groups = product["keywordGroups"]
        result.append([
            product["canonicalId"],
            product["family"],
            product.get("model") or "—",
            join(groups["primary"]),
            join(groups["secondary"]),
            join(groups["longTail"]),
            join(groups["buyerIntent"]),
        ])
    return result


def build_markdown(summary, records, products, issues):
    raw_status_counts = Counter(status for record in records for status in record["statuses"])
    canonical_status_counts = Counter(product["productStatus"] for product in products)
    category_canonical = Counter(product["category"] for product in products)
    issue_counts = Counter(issue["status"] for issue in issues)
    lines = [
        "# COWIN MACHINE Product Research Report",
        "",
        "**Scope:** research-only audit of the current public COWIN MACHINE product-detail URLs. No product page, deployment, or source product data was changed.",
        "",
        f"**Audit timestamp:** {summary['auditedAt']}",
        "",
        "## 1. Executive Summary",
        "",
        f"The public sitemap exposed **{summary['totalRawRecords']}** product-detail records across six categories. These records consolidate into **{summary['canonicalProductFamilies']}** canonical product families under a conservative exact-model or exact-title grouping rule. **{summary['duplicateRawRecords']}** raw records belong to duplicate or near-duplicate groups, and **{summary['modelsNeedingOwnerConfirmation']}** canonical families require owner confirmation before technical page rewrites.",
        "",
        "Products marked `verified-model` can move into fact-checked research and rewrite planning, subject to later evidence verification. Families marked `requires-owner-confirmation`, `duplicate`, or `generic-series` must not receive new technical claims until COWIN MACHINE supplies the missing approved model, specification, and/or identity evidence.",
        "",
        "## 2. Audit Scope and Methodology",
        "",
        "The audit read the public sitemap and then fetched every URL matching `/products/[category]/[slug]`. Each page was parsed for its URL, category slug, product slug, H1, visible model reference, first product description paragraph, visible technical-specification table, local product-image path, related-equipment links, and sitemap last-modified value. The audit does not treat a visible page record as a confirmed SKU. Exact current-model or exact-title groupings were used to identify consolidation candidates; no presumed correction was made for missing, generic, or inconsistent data.",
        "",
        "## 3. Current Website Product Inventory Summary",
        "",
        "| Category | Raw records | Canonical families | Duplicate raw records |",
        "|---|---:|---:|---:|",
    ]
    for category, count in summary["categoryCounts"].items():
        duplicates = sum(1 for record in records if record["category"] == category and "duplicate-or-near-duplicate" in record["statuses"])
        lines.append(f"| {CATEGORY_LABELS[category]} | {count} | {category_canonical[category]} | {duplicates} |")
    lines.extend(["", "Raw status occurrences:", ""])
    for status, count in sorted(raw_status_counts.items()):
        lines.append(f"- `{status}`: {count}")

    lines.extend(["", "## 4. Full URL-Level Product Inventory", "", "The following list includes all 202 audited URLs. The complete captured description, visible specification map, image path, related links, and timestamps are retained in `data/product-audit/raw-product-inventory.json`.", "", "| # | Category | H1 | Model reference | Specs | Image | Related links | Status | URL |", "|---:|---|---|---|---:|---|---:|---|---|"])
    for index, record in enumerate(records, 1):
        lines.append(f"| {index} | {CATEGORY_LABELS[record['category']]} | {text(record.get('currentH1')) or 'not captured'} | {text(record.get('modelReference')) or 'not captured'} | {len(record['visibleSpecifications'])} | {'yes' if record.get('imagePath') else 'no'} | {len(record['relatedLinks'])} | {join(record['statuses'], '; ')} | {record['url']} |")

    lines.extend(["", "## 5. Canonical Product Families and Series", "", "| Canonical ID | Category | Family | Model | Raw records | Status | Current URL mapping |", "|---|---|---|---|---:|---|---|"])
    for product in products:
        lines.append(f"| {product['canonicalId']} | {CATEGORY_LABELS[product['category']]} | {product['family']} | {product.get('model') or '—'} | {len(product['rawRecordIds'])} | {product['productStatus']} | {join(product['currentUrls'], '<br>')} |")

    lines.extend(["", "## 6. Full Product Keyword Master", "", "| Canonical ID | Product family / model | Primary | Secondary | Long-tail | Buyer-intent |", "|---|---|---|---|---|---|"])
    for product in products:
        groups = product['keywordGroups']
        lines.append(f"| {product['canonicalId']} | {product['family']} {product.get('model') or ''} | {join(groups['primary'])} | {join(groups['secondary'])} | {join(groups['longTail'])} | {join(groups['buyerIntent'])} |")

    lines.extend(["", "## 7. Duplicate, Generic and Misclassified Product Data Issues", "", "| Issue status | Occurrences | Handling rule |", "|---|---:|---|"])
    guidance = {
        "duplicate-or-near-duplicate": "Confirm whether URLs are unique SKUs; consolidate only with owner approval.",
        "family-or-series": "Treat as a family/series until an approved model reference is supplied.",
        "generic-title-needs-identification": "Do not publish model-specific technical content.",
        "missing-specification": "Obtain an approved model-specific specification table before rewriting.",
        "missing-image": "Keep a placeholder until an authorized product image is available.",
        "missing-source-evidence": "Request owner or manufacturer evidence for page identity.",
        "possible-category-mismatch": "Require taxonomy review before moving the page.",
    }
    for status, count in sorted(issue_counts.items()):
        lines.append(f"| {status} | {count} | {guidance.get(status, 'Owner review required.')} |")

    lines.extend(["", "## 8. Missing Data Checklist", "", "- Approved manufacturer model / approved family designation for each generic or series page.", "- Model-specific technical specification table and test conditions where applicable.", "- Authorized, watermark-free original image with file-rights confirmation.", "- Approved product documentation / catalogue reference where a technical claim is required.", "- Owner decision for every duplicate or near-duplicate group before URL consolidation or rewrite.", "", "## 9. Recommended Product Taxonomy", "", "Retain the six current top-level families for audit continuity, then organize by verified equipment function rather than by generic promotional title: compressed air equipment; generator systems; drilling equipment; drilling tools & consumables; solar & mobile light towers; magnetic separators. Within each category, add product families only after COWIN MACHINE confirms a stable model or series definition.", "", "## 10. Recommended Data Fields by Product Category", ""])
    for category, field_list in RECOMMENDED_FIELDS.items():
        lines.append(f"- **{CATEGORY_LABELS[category]}:** {field_list}")

    lines.extend(["", "## 11. Priority List for Product Detail Page Rewrite", "", "1. Prioritize the 68 canonical families currently marked `verified-model` for evidence review and independent rewrite briefs.", "2. Resolve the 20 duplicate canonical groups before producing page-level content, titles, or canonical decisions.", "3. Obtain owner confirmation for the 21 families with model-specific specification gaps and the 26 generic-series families.", "4. Do not publish new technical, certification, application-fit, performance, or packaging claims for any unresolved family.", "", "## 12. Recommended Next Actions", "", "1. Ask the owner to approve the canonical mapping and duplicate-consolidation decisions.", "2. Collect model-specific datasheets, authorized images, and verified technical specifications for the 67 families requiring owner confirmation.", "3. Build approved product fact cards from those sources before any page rewrite work starts.", "4. Prepare independent new-brand content briefs from the canonical master; do not copy source-page metadata, H1s, FAQs, or visual assets without rights clearance.", "5. After fact cards are approved, perform a separate content and SEO implementation phase with a pre-publication evidence review.", ""])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(summary, records, products, issues):
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_page(doc.sections[0])
    build_styles(doc)
    add_header_footer(doc.sections[0])

    # Editorial-cover opening for a dense reference report.
    doc.add_paragraph().paragraph_format.space_after = Pt(82)
    add_para(doc, "PRODUCT DATA AUDIT", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color="7A5A00")
    add_para(doc, "COWIN MACHINE", align=WD_ALIGN_PARAGRAPH.CENTER, size=31, color=INK, bold_prefix="COWIN MACHINE")
    add_para(doc, "Product Research Report", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, color=DARK_BLUE)
    add_para(doc, "Public website inventory, canonical-family mapping, keyword master and data-quality review", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    add_para(doc, f"Audit timestamp: {summary['auditedAt']}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=INK, bold_prefix="Audit timestamp: ")
    add_para(doc, "Research-only local deliverable. No website page, deployment, or source product data was changed.", align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()

    raw_status_counts = Counter(status for record in records for status in record["statuses"])
    canonical_status_counts = Counter(product["productStatus"] for product in products)
    category_canonical = Counter(product["category"] for product in products)
    issue_counts = Counter(issue["status"] for issue in issues)

    add_section_heading(doc, 1, "Executive Summary")
    add_para(doc, f"The current public sitemap contains {summary['totalRawRecords']} product-detail page records across six product categories. A conservative audit grouped them into {summary['canonicalProductFamilies']} canonical product families. {summary['duplicateRawRecords']} raw records belong to duplicate or near-duplicate groups, while {summary['modelsNeedingOwnerConfirmation']} canonical families require owner confirmation before technical page rewrite work.")
    add_table(doc, ["Metric", "Result", "Interpretation"], [
        ["URL-level page records", str(summary["totalRawRecords"]), "Current website records; not assumed to be confirmed independent SKUs."],
        ["Canonical product families", str(summary["canonicalProductFamilies"]), "Exact-model or exact-title consolidation candidates with URL mapping preserved."],
        ["Duplicate raw records", str(summary["duplicateRawRecords"]), "Require consolidation/identity review before rewrite or canonical changes."],
        ["Families needing owner confirmation", str(summary["modelsNeedingOwnerConfirmation"]), "Includes generic-series, duplicate, and specification-gap families."],
    ], [2100, 1100, 6160], font_size=9)
    add_para(doc, "Safe conclusion: pages and families marked verified-model may proceed to fact-checking and rewrite research, not automatic publishing. All requires-owner-confirmation, duplicate, and generic-series records must remain free of new technical claims until approved evidence is supplied.", fill=WARNING_FILL, bold_prefix="Safe conclusion: ")

    add_section_heading(doc, 2, "Audit Scope and Methodology")
    add_para(doc, "The audit read the current public sitemap, filtered every URL matching /products/[category]/[slug], and fetched all 202 product-detail pages. It captured page URL, category and product slugs, H1, visible model reference, the first product description paragraph, visible technical-specification table, local product-image path, related-equipment links, and sitemap last-modified value.")
    add_bullets(doc, [
        "A page record was not automatically treated as a unique SKU.",
        "Visible H1 model codes were extracted only when they matched the supplied category-series direction and included a digit; bare series labels were retained as unconfirmed families.",
        "Duplicate detection is conservative: exact current model or exact current title within the same category. Near-duplicate candidates remain owner-review items rather than assumed merges.",
        "This report records visible evidence only. It does not infer missing specifications, certification, image rights, application fit, or correct category placement.",
    ])

    add_section_heading(doc, 3, "Current Website Product Inventory Summary")
    summary_rows = []
    for category, count in summary["categoryCounts"].items():
        duplicates = sum(1 for record in records if record["category"] == category and "duplicate-or-near-duplicate" in record["statuses"])
        verified = sum(1 for record in records if record["category"] == category and "verified-model" in record["statuses"])
        summary_rows.append([CATEGORY_LABELS[category], str(count), str(category_canonical[category]), str(duplicates), str(verified)])
    add_table(doc, ["Category", "Raw\nrecords", "Canonical\nfamilies", "Duplicate\nraw", "Verified\nmodel pages"], summary_rows, [3100, 1250, 1450, 1400, 1360], font_size=8.5)
    status_rows = [[status, str(count), "Raw page-status occurrence"] for status, count in sorted(raw_status_counts.items())]
    add_table(doc, ["Raw record status", "Count", "Meaning"], status_rows, [3000, 1100, 5260], font_size=8.5)
    canonical_rows_data = [[status, str(count), "Canonical-family status"] for status, count in sorted(canonical_status_counts.items())]
    add_table(doc, ["Canonical status", "Count", "Meaning"], canonical_rows_data, [3000, 1100, 5260], font_size=8.5)

    # Landscape detailed inventory with all URL-level records.
    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(landscape, landscape=True)
    add_header_footer(landscape)
    add_section_heading(doc, 4, "Full URL-Level Product Inventory")
    add_para(doc, "All 202 current product-detail URLs are listed below. The companion raw JSON contains the complete captured description text, specification maps, image paths, related links, and timestamps for each corresponding raw record.", style="Report Small")
    add_table(doc, ["#", "Category", "Current H1 / model reference", "Captured completeness", "Status", "Current URL path"], inventory_rows(records), [360, 1450, 3100, 2650, 2100, 3300], font_size=6.8)

    add_section_heading(doc, 5, "Canonical Product Families and Series")
    add_para(doc, "Duplicate and series mappings are retained rather than removed. A canonical entry represents an audit consolidation candidate only; it is not a product-identity approval.", style="Report Small")
    add_table(doc, ["Canonical ID", "Category", "Family", "Model", "Raw\nrecords", "Status", "Current URL mapping"], canonical_rows(products), [2100, 1500, 2700, 1250, 800, 1450, 4760], font_size=6.8)

    add_section_heading(doc, 6, "Full Product Keyword Master")
    add_para(doc, "Keyword groups are research terms organized from supplied category directions and current family/model names. They are not claims, rankings, or publication-ready copy.", style="Report Small")
    add_table(doc, ["Canonical ID", "Family", "Model", "Primary", "Secondary", "Long-tail", "Buyer-intent"], keyword_rows(products), [1450, 1950, 950, 1750, 2500, 2200, 1710], font_size=6.5)

    portrait = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page(portrait, landscape=False)
    add_header_footer(portrait)
    add_section_heading(doc, 7, "Duplicate, Generic and Misclassified Product Data Issues")
    add_para(doc, "The detailed issue register is retained in data/product-audit/product-data-quality-issues.csv. The following aggregate view shows all identified issue classes and their safe handling rules.")
    guidance = {
        "duplicate-or-near-duplicate": "Confirm whether the URL is a distinct SKU; consolidate only after owner approval.",
        "family-or-series": "Treat as a family/series until an approved model reference is supplied.",
        "generic-title-needs-identification": "Do not write model-specific technical content.",
        "missing-specification": "Obtain an approved model-specific technical table before a technical rewrite.",
        "missing-image": "Use a placeholder until an authorized image is supplied.",
        "missing-source-evidence": "Request owner or manufacturer evidence for page identity.",
        "possible-category-mismatch": "Require taxonomy review before moving a record.",
    }
    add_table(doc, ["Issue status", "Occurrences", "Safe handling rule"], [[status, str(count), guidance.get(status, "Owner review required.")] for status, count in sorted(issue_counts.items())], [2800, 1200, 5360], font_size=8.5)
    dup_groups = [product for product in products if product["productStatus"] == "duplicate"]
    add_para(doc, f"{len(dup_groups)} canonical groups have more than one mapped raw record. The largest groups and their preserved URL counts are listed below; the complete mapping is in Section 5.")
    add_table(doc, ["Canonical family", "Category", "Model", "Mapped URL records", "Owner decision required"], [[product["family"], CATEGORY_LABELS[product["category"]], product.get("model") or "—", str(len(product["rawRecordIds"])), "Confirm unique SKU vs. consolidated family"] for product in sorted(dup_groups, key=lambda item: len(item["rawRecordIds"]), reverse=True)], [3000, 1900, 1200, 1400, 1860], font_size=7.5)

    add_section_heading(doc, 8, "Missing Data Checklist")
    add_bullets(doc, [
        "Approved manufacturer model or approved family designation for generic and bare-series pages.",
        "Model-specific technical specification table and applicable test conditions.",
        "Authorized, watermark-free product image with file-rights confirmation.",
        "Approved catalogue, manual, or manufacturer documentation reference for each technical claim.",
        "Owner decision for every duplicate or near-duplicate group before product URL, canonical, or content changes.",
    ])
    missing_rows = []
    for category in summary["categoryCounts"]:
        generic = sum(1 for record in records if record["category"] == category and "generic-title-needs-identification" in record["statuses"])
        no_specs = sum(1 for record in records if record["category"] == category and "missing-specification" in record["statuses"])
        no_evidence = sum(1 for record in records if record["category"] == category and "missing-source-evidence" in record["statuses"])
        missing_rows.append([CATEGORY_LABELS[category], str(generic), str(no_specs), str(no_evidence)])
    add_table(doc, ["Category", "Generic / series pages", "Missing visible specs", "Missing source evidence"], missing_rows, [3300, 1900, 2200, 1960], font_size=8.5)

    add_section_heading(doc, 9, "Recommended Product Taxonomy")
    add_para(doc, "Maintain the current six top-level categories for audit continuity. Within each category, use approved equipment-function families and confirmed models rather than generic promotional labels. No category move should be made merely from a model prefix; the current audit did not identify a conclusive visible category mismatch.")
    taxonomy_rows = [[CATEGORY_LABELS[category], "Retain top level", "Create subfamilies only after approved model/series confirmation"] for category in summary["categoryCounts"]]
    add_table(doc, ["Top-level category", "Recommendation", "Taxonomy control"], taxonomy_rows, [3100, 1900, 4360], font_size=8.5)

    add_section_heading(doc, 10, "Recommended Data Fields by Product Category")
    add_table(doc, ["Category", "Required fact-card fields before technical rewrite"], [[CATEGORY_LABELS[key], value] for key, value in RECOMMENDED_FIELDS.items()], [2900, 6460], font_size=8.5)

    add_section_heading(doc, 11, "Priority List for Product Detail Page Rewrite")
    add_table(doc, ["Priority", "Audience", "Reason", "Publication constraint"], [
        ["1", "68 verified-model canonical families", "Most complete visible model identity in the current audit.", "Research and fact-check first; publish only after verified source material is approved."],
        ["2", "20 duplicate canonical groups", "Multiple current URL records map to the same current exact model/title key.", "Resolve identity and consolidation decision before content, canonical, or URL work."],
        ["3", "21 specification-gap families", "A current model reference is present but fewer than three visible spec fields were captured.", "No technical rewrite until model-specific specifications are approved."],
        ["4", "26 generic-series families", "No approved model-specific reference was captured from the current page.", "Do not publish model-specific or performance content."],
    ], [650, 2300, 3000, 3410], font_size=8.2)

    add_section_heading(doc, 12, "Recommended Next Actions")
    add_numbered(doc, [
        "Review and approve the canonical mapping and duplicate-consolidation decisions with the site owner.",
        "Collect model-specific datasheets, authorized images, and approved technical specifications for the 67 canonical families that need owner confirmation.",
        "Create signed-off product fact cards before drafting any new product-page copy, schema, filters, or internal links.",
        "Prepare independent rewrite briefs from the approved canonical master. Do not reuse visible page metadata, H1s, FAQs, or image assets without documented authorization.",
        "Run a separate implementation, SEO, and pre-publication evidence review only after the fact cards are complete.",
    ])
    add_para(doc, "Publication guardrail: until the missing evidence is supplied, unresolved pages may remain as non-technical placeholders but must not gain technical performance, certification, packaging, application-fit, or availability claims.", fill=WARNING_FILL, bold_prefix="Publication guardrail: ")

    doc.core_properties.title = "COWIN MACHINE Product Research Report"
    doc.core_properties.subject = "Research-only product data audit"
    doc.core_properties.author = "COWIN MACHINE product audit workflow"
    doc.core_properties.comments = "Generated from the local public-site audit data."
    doc.save(DOCX_PATH)


def main():
    summary, records, products, issues = read_data()
    if len(records) != 202:
        raise SystemExit(f"Expected 202 records, found {len(records)}")
    build_markdown(summary, records, products, issues)
    build_docx(summary, records, products, issues)
    print(json.dumps({
        "docx": str(DOCX_PATH),
        "markdown": str(MD_PATH),
        "rawRecords": len(records),
        "canonicalProducts": len(products),
        "issues": len(issues),
        "generatedAt": datetime.now(timezone.utc).isoformat(),
    }, indent=2))


if __name__ == "__main__":
    main()
